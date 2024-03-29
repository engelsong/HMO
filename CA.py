#!/usr/bin/python3
# encoding: utf-8
"""
@version: python3.6
@author: ‘Song‘
@software: HMO
@file: CA.py
@time: 9:50
"""

import string
import re
import cmath
import os
from datetime import datetime
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font, PatternFill
from openpyxl.workbook.properties import CalcProperties
from openpyxl.worksheet.properties import PageSetupProperties
from openpyxl.worksheet.page import PageMargins
from os import linesep, popen, listdir
from openpyxl import load_workbook
from docx.enum import text
# from openpyxl.formatting.rule import CellIsRule
from docx import oxml
from docx.shared import Pt


class Project(object):
    """通过Word文档建立项目对象保存项目信息"""

    def __init__(self, document_name):
        self.name = None
        self.code = None
        self.date = None
        self.destination = None
        self.trans = None
        self.trans_time = None
        self.totalsum = 0
        self.is_lowprice = False  # 是否为低价法
        self.sec_comlist = False  # 是否有供货清单二
        self.is_tech = False  # 是否有技术服务
        self.is_qa = False  # 是否有售后
        self.is_cc = False  # 是否来华培训
        self.techinfo = []  # 存放技术服务信息，格式为[人数，天数，[伙食费，住宿费，公杂费]]
        self.training_days = 0  # 来华培训天数
        self.training_num = 0  # 来华培训人数
        self.qc = []  # 法检物资序号
        self.commodities = {}  # 存放物资信息字典
        self.commodities2 = {}  # 存放供货清单二物资
        document = Document(document_name)
        table1 = document.tables[0]
        table2 = document.tables[1]  # 读取两个表格
        project_info = []
        for cell in table1.column_cells(1):
            project_info.append(cell.text)
        table2_length = len(table2.rows)
        for index in range(1, table2_length):  # 从第2行开始读取表格
            temp = []
            row_now = table2.row_cells(index)
            length_row = len(row_now)
            for i in range(1, length_row):  # 将每行信息放入暂存数组
                temp.append(row_now[i].text)
            temp.append(row_now[0].text)  # 把物资编号放在最后一位
            self.commodities[index] = temp
        self.name, self.code, self.date, self.destination, self.trans, self.trans_time = project_info[0:6]
        self.totalsum = int(project_info[6])

        if project_info[7] in 'yY':
            self.is_lowprice = True
        if project_info[8] in 'yY':
            self.sec_comlist = True
            table3 = document.tables[2]
            self.commodities2 = {}  # 存放供货清单二物资
            # 读取供货清单二
            table3_length = len(table3.rows)
            for index in range(1, table3_length):  # 从第2行开始读取表格
                temp = []
                row_now = table3.row_cells(index)
                length_row = len(row_now)
                for i in range(1, length_row - 1):  # 将每行信息放入暂存数组
                    if i == 6:
                        amount = ''
                        the_unit = ''
                        for d in row_now[i].text:
                            if d.isdigit():
                                amount += d
                        the_unit = row_now[i].text.replace(amount, '')
                        temp.append(amount)
                        temp.append(the_unit)
                    else:
                        temp.append(row_now[i].text)
                price = ''
                for d in row_now[length_row - 1].text:
                    if d.isdigit() or d == '.':
                        price += d
                temp.append(float(price))  # 将金额转换为float
                temp.append(row_now[0].text)  # 把物资编号放在最后一位
                self.commodities2[index] = temp

        if project_info[9] in 'yY':
            self.is_tech = True
            self.techinfo += list(map(int, project_info[10:12]))
        if project_info[12] in 'yY':
            self.is_qa = True
        if project_info[13] in 'yY':
            self.is_cc = True
            self.training_days = int(project_info[15])  # 读取来华陪训天数
            self.training_num = int(project_info[14])  # 读取来华培训人数
        if project_info[-1] != '':
            if project_info[-1] not in 'Nn':
                self.qc += list(map(int, project_info[-1].split()))
                self.qc.sort()

    def show_info(self):
        print('项目名称:', self.name)
        print('项目代码:', self.code)
        print('开标日期:', self.date)
        print('目的地:', self.destination)
        print('运输方式:', self.trans)
        print('运输时间:', self.trans_time)
        print('对外货值：', self.totalsum)
        print('是否为低价法', '是' if self.is_lowprice is True else '否')
        print('是否有供货清单二', '是' if self.sec_comlist is True else '否')
        print('是否有技术服务:', '是' if self.is_tech is True else '否')
        print('是否有售后服务:', '是' if self.is_qa is True else '否')
        print('是否有来华培训', '是' if self.is_cc is True else '否')
        if self.is_tech:
            print('技术服务人数:', self.techinfo[0])
            print('技术服务天数:', self.techinfo[1])
        if self.is_cc:
            print('来华培训人数：', self.training_num)
            print('来华培训天数：', self.training_days)
        if len(self.qc) > 0:
            print('法检物资：', self.qc)

    def show_commodity(self):
        temp_list = sorted(list(self.commodities.keys()))
        for i in temp_list:
            print(i)
            for j in self.commodities[i]:
                print(j)

    def show_commodity2(self):
        temp_list = sorted(list(self.commodities2.keys()))
        for i in temp_list:
            print(self.commodities2[i])
            # for j in self.commodities2[i]:
            #     print(j)


class Quotation(object):
    """通过project实例创建报价表"""

    margin = PageMargins(left=0.25, right=0.25, top=0.75, bottom=0.75, header=0.30, footer=0.30)
    fitsetup = PageSetupProperties(fitToPage=True)

    def __init__(self, project):
        self.project = project
        self.wb = Workbook()
        self.ws_input = None
        self.ws_cost = None
        self.ws_examination = None
        self.ws_lawexam = None
        self.ws_techserve = None
        self.ws_selection = None
        self.ws_itemized_quotation = None
        self.ws_tax_refund = None
        self.ws_summed_quotation = None
        self.ws_general = None
        self.ws_training = None
        self.ws_isolist = None
        self.ws_conservlist = None
        self.ws_eplist = None
        self.ws_lob = None
        self.ws_self_exam = None
        self.ws_3rd_exam = None

    def create_all(self):
        self.create_input()
        self.create_cost()
        self.create_selection()
        self.create_lob()
        self.create_examination()
        if len(self.project.qc) > 0:
            self.create_lawexam()
        if self.project.is_cc:
            self.create_training()
        if self.project.is_tech:
            self.create_techserve()
        self.create_tax_refund()
        self.create_itemized_quotation()
        # self.create_summed_quotation()
        self.create_general()
        self.create_isolist()
        self.create_conservlist()
        self.create_eplist()
        self.create_self_exam()
        self.create_3rd_exam()
        self.wb.calculation = CalcProperties(iterate=True)
        self.wb.save('投标报价表-{}.xlsx'.format(self.project.name))

    def create_general(self):
        """创建总报价表"""
        self.ws_general = self.wb.create_sheet('1.投标报价总表', 4)
        colum_title = ['序号', '费用项目', '合计金额', '备注']
        title_width = [10, 35, 25, 20]
        colum_number = len(colum_title)
        row_number = 7
        if self.project.is_tech:
            row_number += 1
        if self.project.is_cc:
            row_number += 1

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=14)
        normal_font = Font(name='宋体', size=14)
        title_font = Font(name='黑体', bold=True, size=20)

        # 初始化表格
        for i in range(colum_number):
            for j in range(row_number):
                cell_now = self.ws_general.cell(row=j + 1, column=i + 1)
                if j > 1:
                    cell_now.border = full_border
                if j < 3 and j != 1:
                    cell_now.font = bold_font
                else:
                    cell_now.font = normal_font
                if i == 1 and j > 2:
                    cell_now.alignment = Alignment(
                        horizontal='left', vertical='center', wrap_text=True)
                elif i == 2 and j > 2:
                    cell_now.alignment = Alignment(
                        horizontal='right', vertical='center', wrap_text=True)
                    cell_now.number_format = '¥#,##0.00'
                else:
                    if j != 1:
                        cell_now.alignment = ctr_alignment
        for i in range(len(title_width)):  # 修改列宽
            self.ws_general.column_dimensions[
                self.ws_general.cell(row=4, column=i + 1).column_letter].width = title_width[i]
        self.ws_general.row_dimensions[3].height = 40

        # 创建标题行
        self.ws_general.merge_cells('A2:D2')
        self.ws_general.merge_cells('A1:D1')
        self.ws_general['A1'].font = title_font
        self.ws_general['A1'].alignment = ctr_alignment
        self.ws_general['A1'] = '一.投标报价总表'
        self.ws_general.row_dimensions[1].height = 50
        self.ws_general['A2'] = '报价单位：人民币元'

        # 填写表头
        index = 0
        for i in self.ws_general['A3':'D3'][0]:
            # print(index+1, i)
            if colum_title[index] != '':
                i.value = colum_title[index]
                i.font = bold_font
            index += 1

        # 填写数据
        self.ws_general['A4'] = '一'
        self.ws_general['B4'] = "全部物资价格{}（含商品购买价款、国内运杂费、包装费、 保管费、物资检验费、运输保险费、" \
                                "国外运费、资金占用成本、合理利润、税金）".format(linesep)
        self.ws_general['C4'] = "='2.物资对内分项报价表'!M{}".format(
            len(self.project.commodities) + 4)
        self.ws_general['A5'] = "二"

        if self.project.is_tech:
            self.ws_general['C5'] = "='4.技术服务费报价表'!H14"
            self.ws_general['B5'] = '技术服务费'
            if self.project.is_cc:
                self.ws_general['C6'] = "='5.来华培训费报价表'!G17"
                self.ws_general['B6'] = '来华培训费'
        elif self.project.is_cc:
            self.ws_general['C5'] = "='5.来华培训费报价表'!G17"
            self.ws_general['B5'] = '来华培训费'

        no_seq = ['二', '三', '四', '五']
        for i in range(5, row_number):
            self.ws_general['A{}'.format(i)] = no_seq[i - 5]
        self.ws_general["B{}".format(row_number - 2)] = "其他费用{}（含须中方承担的其他费用、管理费、风险预涨费、防恐措施费、" \
                                                        "大型机电设备跟踪反馈工作费用等）".format(linesep)
        self.ws_general['C{}'.format(row_number - 2)] = "=费用输入!J17"
        self.ws_general['B{}'.format(row_number - 1)] = '《供货清单（一）》中各项物资增值税退抵税额'
        self.ws_general['C{}'.format(row_number - 1)] = \
            "='3.各项物资增值税退抵税额表'!F{}".format(len(self.project.commodities) + 4)
        self.ws_general['B{}'.format(row_number)] = '合计'
        self.ws_general['C{}'.format(row_number)] = "=SUM(C4:C{})-C{}".format(
            row_number - 2, row_number - 1)
        self.ws_general['C{}'.format(row_number)].font = bold_font
        for i in range(7, row_number, -1):
            self.ws_general.delete_rows(i)

        # 打印设置
        self.ws_general.print_options.horizontalCentered = True
        self.ws_general.print_area = 'A1:D{}'.format(row_number)
        self.ws_general.page_setup.fitToWidth = 1
        self.ws_general.sheet_properties.pageSetUpPr = Quotation.fitsetup
        self.ws_general.page_margins = Quotation.margin

    def create_input(self):
        """创建物资输入表"""
        self.ws_input = self.wb.create_sheet('物资输入', 0)
        colum_title = ['序号', '品名', 'HS编码', '数量', '', '品牌', '型号', '规格', '单价',
                       '总价', '生产厂商', '供货商', '生产或供货地', '联系人',
                       '联系电话', '出厂日期', '出口港', '检验标准', '检验机构', '交货期', '交货地点',
                       '三体系', '节能', '环保', '备注']
        title_width = [6, 14, 12, 3, 5, 10, 10, 30, 14, 16, 10, 10, 10, 10, 15, 10, 10, 15,
                       16, 10, 10, 15, 15, 15, 6]

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=12)
        normal_font = Font(name='宋体', size=10)

        # 初始化表格
        colum_number = len(colum_title)
        row_number_total = len(self.project.commodities) + len(self.project.commodities2) + 1
        for i in range(colum_number):
            for j in range(row_number_total):
                cell_now = self.ws_input.cell(row=j + 1, column=i + 1)
                cell_now.border = full_border
                cell_now.font = normal_font
                cell_now.alignment = ctr_alignment

        for i in range(len(title_width)):  # 修改列宽
            self.ws_input.column_dimensions[
                self.ws_input.cell(row=1, column=i + 1).column_letter].width = title_width[i]

        # 填写表头
        index = 0
        for i in self.ws_input['A1':'Y1'][0]:
            i.value = colum_title[index]
            i.font = bold_font
            i.alignment = ctr_alignment
            index += 1

        # 填写供货清单一的物资数据
        row_number = len(self.project.commodities) + 1
        relate_coord = [('B', 0), ('C', 1), ('D', 2), ('R', 5)]
        for num in range(2, row_number + 1):
            if self.project.commodities[num - 1][-1] == '':
                self.ws_input['A{}'.format(num)] = num - 1
            else:
                self.ws_input['A{}'.format(
                    num)] = self.project.commodities[num - 1][-1]  # 填写物资序号
            self.ws_input['I{}'.format(num)].number_format = '¥#,##0.00'
            self.ws_input['I{}'.format(num)] = 1
            # self.ws_input['H{}'.format(num)].value = 1
            self.ws_input['J{}'.format(num)].number_format = '¥#,##0.00'
            self.ws_input['Y{}'.format(num)] = '-'
            for rela in relate_coord:
                self.ws_input['{}{}'.format(
                    rela[0], num)] = self.project.commodities[num - 1][rela[1]]
            else:
                self.ws_input['E{}'.format(num)].number_format = '0'
                self.ws_input['E{}'.format(num)] = self.project.commodities[num - 1][3]
            self.ws_input['J{}'.format(num)] = '=E{}*I{}'.format(num, num)

        # 填写供货清单二的物资数据
        relate_coord2 = [('B', 0), ('C', 1), ('K', 2), ('L', 2), ('F', 3), ('H', 4), ('D', 6), ('E', 5), ('J', -2),
                         ('A', -1), ('R', -4)]
        for num in range(row_number + 1, row_number_total + 1):
            num_now = num - row_number
            self.ws_input['I{}'.format(num)].number_format = '¥#,##0.00'
            # self.ws_input['H{}'.format(num)].value = 1
            self.ws_input['J{}'.format(num)].number_format = '¥#,##0.00'
            self.ws_input['Y{}'.format(num)] = '-'
            self.ws_input['E{}'.format(num)].number_format = '0'
            for rela in relate_coord2:
                self.ws_input['{}{}'.format(
                    rela[0], num)] = self.project.commodities2[num_now][rela[1]]
        # self.ws_input.merge_cells('A{}:Y{}'.format(row_number + 1, row_number + 1))
        self.ws_input.merge_cells('D1:E1')

    def create_cost(self):
        """创建费用输入表格"""
        self.ws_cost = self.wb.create_sheet('费用输入', 1)
        colum_title = [
            '海运',
            '单价',
            '',
            '数量',
            '总金额',
            '陆运',
            '单价',
            '',
            '数量',
            '总金额']
        title_width = [14, 12, 12, 8, 16, 14, 16, 12, 8, 16]

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=12)
        normal_font = Font(name='宋体', size=12)
        yellow_fill = PatternFill(
            fill_type='solid',
            start_color='FFFF00',
            end_color='FFFF00')

        # 初始化表格
        colum_number = len(colum_title)
        row_number = 18
        for i in range(colum_number):
            for j in range(row_number):
                cell_now = self.ws_cost.cell(row=j + 1, column=i + 1)
                cell_now.border = full_border
                cell_now.font = normal_font
                cell_now.alignment = ctr_alignment
        for i in range(len(title_width)):  # 修改列宽
            self.ws_cost.column_dimensions[
                self.ws_cost.cell(row=4, column=i + 1).column_letter].width = title_width[i]
        for i in range(1, row_number + 1):  # 修改行高
            self.ws_cost.row_dimensions[
                self.ws_cost.cell(row=i, column=1).row].height = 24

        # 填写表头
        index = 0
        for i in self.ws_cost['A1':'J1'][0]:
            # print(index+1, i)
            if colum_title[index] != '':
                i.value = colum_title[index]
                i.font = bold_font
            index += 1

        # 按列从左到右填写表格
        self.ws_cost['A2'] = '全程运费'
        self.ws_cost['A5'] = '港杂费'
        self.ws_cost['A8'] = '仓库装箱费'
        self.ws_cost['A11'] = '运抵报告费'
        self.ws_cost['A14'] = '加固费'
        self.ws_cost['A15'] = '舱单费'
        self.ws_cost['A16'] = '文件费'
        self.ws_cost['A17'] = '报关费'
        self.ws_cost['A18'] = '苫布费'
        index = 2
        for i in ['20GP', '40GP/HQ', '40FR'] * 4:
            self.ws_cost['B{}'.format(index)] = i
            index += 1
        for i in range(2, 14):
            if i < 5:  # 美元部分格式及E列公式
                self.ws_cost['C{}'.format(i)].number_format = '$#,##0.00'
                self.ws_cost['E{}'.format(i)].number_format = '$#,##0.00'
                self.ws_cost['H{}'.format(i)].number_format = '$#,##0.00'
                self.ws_cost['E{}'.format(i)] = '=C{0}*D{0}*F10'.format(i)
            else:
                self.ws_cost['C{}'.format(i)].number_format = '¥#,##0.00'
            self.ws_cost['C{}'.format(i)] = 0
        self.ws_cost['C14'] = 1000
        self.ws_cost['C14'].number_format = '0"元/箱"'
        self.ws_cost['C15'] = 100
        self.ws_cost['C15'].number_format = '0"元/票"'
        self.ws_cost['C16'] = 500
        self.ws_cost['C16'].number_format = '0"元/票"'
        self.ws_cost['C17'] = 300
        self.ws_cost['C17'].number_format = '0"元/票"'
        self.ws_cost['C18'] = 0
        self.ws_cost['C18'].number_format = '0"元/柜"'
        for i in range(2, 19):  # D列格式及E列格式
            self.ws_cost['D{}'.format(i)] = 0
            self.ws_cost['D{}'.format(i)].number_format = '0'
            self.ws_cost['E{}'.format(i)].number_format = '¥#,##0.00'
            if i > 4:
                self.ws_cost['E{}'.format(
                    i)] = '=C{0}*D{0}'.format(i)  # E 列公式生成
                self.ws_cost['D{}'.format(i)].number_format = '0'
        self.ws_cost['F2'] = '内陆运费'
        self.ws_cost['F5'] = '电子跟踪单'
        self.ws_cost['F7'] = '货物为车辆'
        self.ws_cost['F9'] = '运费合计'
        self.ws_cost['F9'].font = bold_font
        self.ws_cost['F10'].font = bold_font
        self.ws_cost['F10'].fill = yellow_fill
        self.ws_cost['F10'] = 6.6
        self.ws_cost['F10'].number_format = '"汇率："0.0000'
        self.ws_cost['G2'] = '20GP'
        self.ws_cost['G3'] = '40GP/HQ'
        self.ws_cost['G4'] = '40FR'
        self.ws_cost['G5'] = '不超过5个'
        self.ws_cost['G6'] = '超过5个追加'
        self.ws_cost['G7'] = '不超过5个'
        self.ws_cost['G8'] = '超过5个追加'
        for i in range(2, 9):
            self.ws_cost['H{}'.format(i)] = 0
            self.ws_cost['I{}'.format(i)] = 0
            self.ws_cost['J{}'.format(i)] = '=H{0}*I{0}*F10'.format(i)
            self.ws_cost['J{}'.format(i)].number_format = '¥#,##0.00'
        self.ws_cost['J9'] = '=SUM(E2:E18)+SUM(J2:J8)'
        self.ws_cost['J9'].font = bold_font
        self.ws_cost['J9'].number_format = '¥#,##0.00'

        # 保险费用
        self.ws_cost['F18'] = '保险费用'
        self.ws_cost['F18'].font = bold_font
        self.ws_cost['G18'] = self.project.totalsum
        self.ws_cost['G18'].number_format = '¥#,##0.00'
        self.ws_cost['H18'] = '费率'
        self.ws_cost['I18'] = 0.001
        self.ws_cost['I18'].number_format = '0.00%'
        self.ws_cost['I18'].fill = yellow_fill
        self.ws_cost['J18'].font = bold_font
        self.ws_cost['J18'].number_format = '¥#,##0.00'
        self.ws_cost['J18'] = '=round(G18*1.1*I18,2)'

        # 其他费用
        self.ws_cost['F15'] = '管理费'
        self.ws_cost['F15'].font = bold_font
        self.ws_cost['G15'] = 500
        self.ws_cost['G15'].number_format = '¥#,##0.00'
        self.ws_cost['G15'].fill = yellow_fill

        self.ws_cost['H15'] = '风险预张费'
        self.ws_cost['H15'].font = bold_font
        self.ws_cost['J15'] = 500
        self.ws_cost['J15'].number_format = '¥#,##0.00'
        self.ws_cost['J15'].fill = yellow_fill

        self.ws_cost['F16'] = '大型机电费用'
        self.ws_cost['F16'].font = bold_font
        self.ws_cost['G16'] = 0
        self.ws_cost['G16'].number_format = '¥#,##0.00'
        self.ws_cost['G16'].fill = yellow_fill

        self.ws_cost['H16'] = '防恐'
        self.ws_cost['H16'].font = bold_font
        self.ws_cost['J16'] = 0
        self.ws_cost['J16'].number_format = '¥#,##0.00'
        self.ws_cost['J16'].fill = yellow_fill

        self.ws_cost['F17'] = '其他费用'
        self.ws_cost['F17'].font = bold_font
        self.ws_cost['J17'] = '=SUM(J15:J16,G15:G16)'
        self.ws_cost['J17'].number_format = '¥#,##0.00'
        # self.ws_cost['J17'].fill = yellow_fill

        # 商检费用填写
        self.ws_cost['F11'] = '商检费用'
        self.ws_cost['F11'].font = bold_font
        self.ws_cost['J11'] = 0
        self.ws_cost['J11'].fill = yellow_fill
        # 合理利润填写
        self.ws_cost['F13'] = '合理利润'
        self.ws_cost['F13'].font = bold_font
        self.ws_cost['J13'].number_format = '¥#,##0.00'
        self.ws_cost['J13'] = 10000
        self.ws_cost['J13'].fill = yellow_fill

        # 合并需要合并单元格
        self.ws_cost.merge_cells('B1:C1')
        self.ws_cost.merge_cells('G1:H1')
        self.ws_cost.merge_cells('F9:I9')
        self.ws_cost.merge_cells('F10:J10')
        self.ws_cost.merge_cells('F11:I11')
        self.ws_cost.merge_cells('F12:I12')
        self.ws_cost.merge_cells('F17:I17')
        self.ws_cost.merge_cells('F2:F4')
        self.ws_cost.merge_cells('F5:F6')
        self.ws_cost.merge_cells('F7:F8')
        self.ws_cost.merge_cells('F13:I13')
        for i in range(2, 14, 3):
            self.ws_cost.merge_cells('A{}:A{}'.format(i, i + 2))

    def create_selection(self):
        """创建物资选型一览表"""
        self.ws_selection = self.wb.create_sheet('0.物资选型一览表', 2)
        colum_title = ['序号', '物资名称', '招标要求', '投标产品品牌和型号', '生产企业', '供货企业', '交货期',
                       '交货地点', '备注']
        title_width = [6, 12, 50, 20, 8, 8, 6, 6, 6]

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        left_alignment = Alignment(
            horizontal='left',
            vertical='top',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=12)
        normal_font = Font(name='宋体', size=10)
        title_font = Font(name='黑体', size=20)

        # 初始化表格
        colum_number = len(colum_title)
        row_number = len(self.project.commodities) + 2
        for i in range(colum_number):
            for j in range(row_number):
                cell_now = self.ws_selection.cell(row=j + 1, column=i + 1)
                if j > 0:  # 第一列留下给表头
                    cell_now.border = full_border
                    cell_now.font = normal_font
                    cell_now.alignment = ctr_alignment
        for i in range(len(title_width)):  # 修改列宽
            self.ws_selection.column_dimensions[
                self.ws_selection.cell(row=1, column=i + 1).column_letter].width = title_width[i]

        # 创建标题行
        self.ws_selection.merge_cells('A1:I1')
        self.ws_selection['A1'].font = title_font
        self.ws_selection['A1'].alignment = ctr_alignment
        self.ws_selection['A1'] = '2.各项物资选型一览表'
        self.ws_selection.row_dimensions[1].height = 40

        # 填写表头
        index = 0
        for i in self.ws_selection['A2':'I2'][0]:
            i.value = colum_title[index]
            i.font = bold_font
            i.alignment = ctr_alignment
            index += 1

        # 填入数据
        col_relate = [('A', 'A'), ('B', 'B'), ('E', 'K'), ('F', 'L'), ('G', 'T'), ('H', 'U')]
        for row in range(3, row_number + 1):  # 遍历行
            for col in col_relate:  # 根据对应关系设立公式
                cell_now = self.ws_selection['{}{}'.format(col[0], row)]
                cell_now.value = '=物资输入!{}{}'.format(col[1], row - 1)
            self.ws_selection['C{}'.format(
                row)] = self.project.commodities[row - 2][4]  # 填入招标要求
            self.ws_selection['C{}'.format(row)].alignment = left_alignment
            self.ws_selection['D{}'.format(row)] = \
                '="品牌："&物资输入!F{}&CHAR(10)&"型号："&物资输入!G{}'.format(row-1, row-1)

        # 打印设置
        self.ws_selection.print_options.horizontalCentered = True
        self.ws_selection.print_area = 'A1:I{}'.format(row_number)
        self.ws_selection.page_setup.fitToWidth = 1
        self.ws_selection.page_setup.orientation = "landscape"
        self.ws_selection.page_margins = Quotation.margin

    def create_itemized_quotation(self):
        """生成分项报价表垂直方向"""
        self.ws_itemized_quotation = self.wb.create_sheet('2.物资对内分项报价表', 4)
        colum_title = ['物资', '', '商品购买价款', '国内运杂费', '包装费', '保管费', '物资检验费', '运输保险费', '国外运费',
                       '资金占用成本', '合理利润', '税金',
                       '合计（即《供货清单（一）》各项物资{}总价)'.format(self.project.trans)]
        if self.project.sec_comlist:
            colum_title[-1] = '合计（即《供货清单（一）》各项物资{}总价和《供货清单（二）》' \
                              '各项物资流通费用总价)'.format(self.project.trans)

        title_width = [8, 16, 14, 10, 10, 10, 16, 16, 16, 16, 16, 16, 20]
        colum_number = len(colum_title)
        row_number = len(self.project.commodities) + len(self.project.commodities2) + 6

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        slash_border = Border(diagonal=real_side, diagonalDown=True, left=real_side, right=real_side,
                              top=real_side, bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=12)
        normal_font = Font(name='宋体', size=12)
        normal_white_font = Font(name='宋体', color='FFFFFF', size=12)
        title_font = Font(name='黑体', size=14)
        right_alignment = Alignment(
            horizontal='right',
            vertical='center',
            wrap_text=False)
        left_alignment = Alignment(
            horizontal='left',
            vertical='center',
            wrap_text=False)
        # yellow_fill = PatternFill(
        #     fill_type='solid',
        #     start_color='FFFF00',
        #     end_color='FFFF00')

        # 初始化表格
        for i in range(colum_number):
            for j in range(2, row_number):  # 留出第一二行
                cell_now = self.ws_itemized_quotation.cell(
                    row=j + 1, column=i + 1)
                if j == 2:
                    cell_now.font = bold_font
                    cell_now.alignment = ctr_alignment
                else:
                    cell_now.font = normal_font
                if i > 1 and j < row_number - 1:  # 格式化单元格
                    cell_now.number_format = '#,##0.00'
                    cell_now.alignment = right_alignment
                else:
                    cell_now.alignment = ctr_alignment
                if j < row_number - 2:
                    cell_now.border = full_border
                else:
                    cell_now.alignment = left_alignment
        for i in range(len(title_width)):  # 修改列宽
            self.ws_itemized_quotation.column_dimensions[
                self.ws_itemized_quotation.cell(row=4, column=i + 1).column_letter].width = title_width[i]
        for i in range(4, row_number + 1):  # 修改行高
            self.ws_itemized_quotation.row_dimensions[i].height = 20
        self.ws_itemized_quotation.row_dimensions[3].height = 60

        # 创建标题行
        self.ws_itemized_quotation['A1'].font = title_font
        self.ws_itemized_quotation['A1'].alignment = ctr_alignment
        self.ws_itemized_quotation['A1'] = '二.物资对内分项报价表'
        self.ws_itemized_quotation.row_dimensions[1].height = 32

        # 第二行
        self.ws_itemized_quotation['A2'].font = normal_font
        self.ws_itemized_quotation['A2'].alignment = left_alignment
        self.ws_itemized_quotation['A2'] = '报价单位：人民币元'
        self.ws_itemized_quotation.row_dimensions[2].height = 20

        # 填写表头
        index = 0
        for i in self.ws_itemized_quotation['A3':'M3'][0]:
            i.value = colum_title[index]
            i.font = bold_font
            i.alignment = ctr_alignment
            index += 1

        # 填写数据
        # self.ws_itemized_quotation['A{}'.format(row_number - 1)] = '注：'
        # self.ws_itemized_quotation['B{}'.format(row_number - 1)] = '1.资金占用成本=（商品进价成本+物资检验费+保险费' \
        #                                                            '+国外运费）×3%利率×预计占用3个月/12个月'
        # self.ws_itemized_quotation['B{}'.format(row_number)] = '2.税金=[对内总承包价/（1+增值税税率）]' \
        #                                                        'X增值税税率-当期进项税款'
        # self.ws_itemized_quotation['B{}'.format(row_number - 1)].fill = yellow_fill
        # self.ws_itemized_quotation['B{}'.format(row_number)].fill = yellow_fill
        self.ws_itemized_quotation['B{}'.format(row_number - 2)] = '小计'
        self.ws_itemized_quotation['A4'] = '供货清单（一）'
        if self.project.sec_comlist:
            self.ws_itemized_quotation['A{}'.format(len(self.project.commodities) + 4)] = '供货清单（二）'

        # col_relate = [('A', 'A'), ('B', 'B'), ('C', 'J')]
        row_sum = row_number - 2
        if self.project.sec_comlist:
            row_sum = row_number
        for row in range(4, row_number - 2):
            # for col in col_relate:  # 根据对应关系设立公式
            self.ws_itemized_quotation['C{}'.format(row)] = '=物资输入!J{}'.format(row - 2)
            self.ws_itemized_quotation['B{}'.format(row)] = '=物资输入!A{0}&"."&物资输入!B{0}'.format(row - 2)
            self.ws_itemized_quotation['D{}'.format(row)] = 0
            self.ws_itemized_quotation['E{}'.format(row)] = 0
            self.ws_itemized_quotation['F{}'.format(row)] = 0
            self.ws_itemized_quotation['G{}'.format(
                row)] = '=round(C{0}/C{1}*G{2},2)'.format(row, row_sum, row_number)
            self.ws_itemized_quotation['H{}'.format(
                row)] = '=round(C{0}/C{1}*H{2},2)'.format(row, row_sum, row_number)
            self.ws_itemized_quotation['I{}'.format(
                row)] = '=round(C{0}/C{1}*I{2},2)'.format(row, row_sum, row_number)
            self.ws_itemized_quotation['J{}'.format(
                row)] = '=round(C{0}/C{1}*J{2},2)'.format(row, row_sum, row_number)
            self.ws_itemized_quotation['K{}'.format(
                row)] = '=round(C{0}/C{1}*K{2},2)'.format(row, row_sum, row_number)
            self.ws_itemized_quotation['L{}'.format(
                row)] = '=round(C{0}/C{1}*L{2},2)'.format(row, row_sum, row_number)
            if row < len(self.project.commodities) + 4:
                self.ws_itemized_quotation['M{}'.format(row)] = '=SUM(C{0}:L{0})'.format(row)
            else:
                self.ws_itemized_quotation['M{}'.format(row)] = '=SUM(D{0}:L{0})'.format(row)
        for column in 'DEFGHIJKLM':
            self.ws_itemized_quotation['{}{}'.format(column, row_number - 2)]\
                = '=SUM({0}4:{0}{1})'.format(column, row_number - 3)
        self.ws_itemized_quotation['C{}'.format(row_number - 2)]\
            = '=SUM(C4:C{})'.format(len(self.project.commodities) + 3)

        self.ws_itemized_quotation['G{}'.format(row_number)] = '=费用输入!J11'
        self.ws_itemized_quotation['H{}'.format(row_number)] = '=费用输入!J18'
        self.ws_itemized_quotation['I{}'.format(row_number)] = '=费用输入!J9'
        self.ws_itemized_quotation['K{}'.format(row_number)] = '=费用输入!J13'
        general_row = 7
        if self.project.is_tech:
            general_row += 1
        if self.project.is_cc:
            general_row += 1
        self.ws_itemized_quotation['L{}'.format(row_number)] = \
            "=ROUND((sum(C{0}:I{0})+'1.投标报价总表'!C{1})*0.0003,2)".format(row_number - 2, general_row)
        self.ws_itemized_quotation['M{}'.format(
            row_number)] = '=SUM(C{0}:L{0})'.format(row_number - 2)
        self.ws_itemized_quotation['J{}'.format(row_number)] = \
            '=round(SUM(C{0}:I{0})*3/12*0.0435,2)'.format(row_number - 2)
        if self.project.sec_comlist:
            self.ws_itemized_quotation['C{}'.format(row_number)] = '=SUM(C4:C{})'.format(row_number - 3)
            for row in range(len(self.project.commodities) + 4, row_number - 2):
                self.ws_itemized_quotation['C{}'.format(row)].font = normal_white_font
                self.ws_itemized_quotation['C{}'.format(row)].border = slash_border

        # self.ws_itemized_quotation['J{}'.format(
        #     row_number - 2)].fill = yellow_fill
        # self.ws_itemized_quotation['N{}'.format(
        #     row_number - 2)] = '=SUM(M4:M{})'.format(row_number - 3)

        # # 低价项目针对部分单元格进行修改
        # if self.project.is_lowprice:
        #     for row in range(3, row_number - 3):
        #         self.ws_itemized_quotation['G{}'.format(row)] = 0.01
        #         self.ws_itemized_quotation['H{}'.format(row)] = 0.01
        #         self.ws_itemized_quotation['I{}'.format(row)] = 0.01
        #     self.ws_itemized_quotation['G{}'.format(row_number - 3)] = '=sum(G3:G{})'.format(row_number - 4)
        #     self.ws_itemized_quotation['H{}'.format(row_number - 3)] = '=sum(H3:H{})'.format(row_number - 4)
        #     self.ws_itemized_quotation['I{}'.format(row_number - 3)] = '=sum(I3:I{})'.format(row_number - 4)
        #     self.ws_itemized_quotation['J{}'.format(
        #         row_number - 3)] = '=round((SUM(C{0}:I{0})*3/12*0.0435)*0.8,2)'.format(row_number - 3)
        #     self.ws_itemized_quotation['K{}'.format(
        #         row_number - 3)] = '=round(IF(C{0}>50000000,(C{0}-50000000)*0.0075+835000,IF(C{0}>20000000,' \
        #                            '(C{0}-20000000)*0.01+535000,IF(C{0}>10000000,(C{0}-10000000)*0.02+335000,' \
        #                            'IF(C{0}>5000000,(C{0}-5000000)*0.03+185000,IF(C{0}>2000000,' \
        #                            '(C{0}-2000000)*0.035+80000,C{0}*0.04)))))*0.8,2)'.format(row_number - 3)
        #     self.ws_itemized_quotation['L{}'.format(row_number - 3)] = \
        #         '=round((M{0}/1.13*0.13-C{0}/1.13*0.13-G{0}/1.06*0.06)*0.9,2)'.format(
        #             row_number - 3)

        # 增加条件格式判断
        # red_fill = PatternFill(
        #     start_color='EE1111',
        #     end_color='EE1111',
        #     fill_type='solid')
        # self.ws_itemized_quotation.conditional_formatting.add('N{}'.format(row_number - 2), CellIsRule(
        #     operator='notEqual', formula=['M{}'.format(row_number - 2)], fill=red_fill))

        # 合并需要合并单元格
        self.ws_itemized_quotation.merge_cells('A1:M1')
        # self.ws_itemized_quotation.merge_cells('B{0}:M{0}'.format(row_number - 1))
        # self.ws_itemized_quotation.merge_cells(
        #     'B{0}:M{0}'.format(row_number - 2))
        # self.ws_itemized_quotation.merge_cells('B{0}:M{0}'.format(row_number))
        self.ws_itemized_quotation.merge_cells('A4:A{}'.format(len(self.project.commodities) + 3))
        if self.project.sec_comlist:
            self.ws_itemized_quotation.merge_cells('A{}:A{}'.format(len(self.project.commodities) + 4, row_number - 3))

        self.ws_itemized_quotation.merge_cells('A3:B3')

        # 打印设置
        self.ws_itemized_quotation.print_options.horizontalCentered = True
        self.ws_itemized_quotation.print_area = 'A1:M{}'.format(row_number - 2)
        self.ws_itemized_quotation.page_setup.fitToWidth = 1
        self.ws_itemized_quotation.page_setup.orientation = "landscape"
        self.ws_itemized_quotation.page_margins = PageMargins(left=0.25, right=0.25, top=0.75, bottom=0.75, header=0.3,
                                                              footer=0.3)

    def create_tax_refund(self):
        """生成退税额表"""
        self.ws_tax_refund = self.wb.create_sheet('3.各项物资增值税退抵税额表', 4)
        colum_title = ['序号', '品名', '投标人向物资生产供货企业支付的商品购买价款（元）',
                       '物资生产供货企业实缴增值税税率（%）', '投标人预期可获得的退抵物资增值税率（%）',
                       '投标人预期可获得的退抵物资增值税额（元）']

        title_width = [8, 16, 30, 22, 25, 30]
        colum_number = len(colum_title)
        row_number = len(self.project.commodities) + 4

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        slash_border = Border(diagonal=real_side, diagonalDown=True, left=real_side, right=real_side,
                              top=real_side, bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=12)
        normal_font = Font(name='宋体', size=12)
        title_font = Font(name='黑体', size=14)
        right_alignment = Alignment(
            horizontal='right',
            vertical='center',
            wrap_text=False)
        left_alignment = Alignment(
            horizontal='left',
            vertical='center',
            wrap_text=False)
        # yellow_fill = PatternFill(
        #     fill_type='solid',
        #     start_color='FFFF00',
        #     end_color='FFFF00')

        # 初始化表格
        for i in range(colum_number):
            for j in range(2, row_number):  # 留出第一二行
                cell_now = self.ws_tax_refund.cell(row=j + 1, column=i + 1)
                cell_now.border = full_border
                if j == 2:  # 分离标题行
                    cell_now.font = bold_font
                    cell_now.alignment = ctr_alignment
                else:
                    cell_now.font = normal_font
                    if i in (2, 5):
                        cell_now.number_format = '#,##0.00'
                        cell_now.alignment = right_alignment
                    else:
                        cell_now.alignment = ctr_alignment

        for i in range(len(title_width)):  # 修改列宽
            self.ws_tax_refund.column_dimensions[
                self.ws_tax_refund.cell(row=4, column=i + 1).column_letter].width = title_width[i]
        for i in range(row_number + 1):  # 修改行高
            self.ws_tax_refund.row_dimensions[i].height = 30
        self.ws_tax_refund.row_dimensions[3].height = 45

        # 创建标题行
        self.ws_tax_refund['A1'].font = title_font
        self.ws_tax_refund['A1'].alignment = ctr_alignment
        self.ws_tax_refund['A1'] = '三.《供货清单（一）》中各项物资增值税退抵税额表'
        self.ws_tax_refund.row_dimensions[1].height = 40

        # 第二行
        self.ws_tax_refund['A2'].font = normal_font
        self.ws_tax_refund['A2'].alignment = left_alignment
        self.ws_tax_refund['A2'] = '报价单位：人民币元'

        # 填写表头
        index = 0
        for i in self.ws_tax_refund['A3':'F3'][0]:
            i.value = colum_title[index]
            i.font = bold_font
            i.alignment = ctr_alignment
            index += 1

        # 填写数据
        self.ws_tax_refund['A{}'.format(row_number)] = '共计'

        for row in range(4, row_number):
            self.ws_tax_refund['A{}'.format(row)] = '=物资输入!A{}'.format(row - 2)
            self.ws_tax_refund['B{}'.format(row)] = '=物资输入!B{}'.format(row - 2)
            self.ws_tax_refund['C{}'.format(row)] = '=物资输入!J{}'.format(row - 2)
            self.ws_tax_refund['D{}'.format(row)] = 13
            self.ws_tax_refund['E{}'.format(row)] = 13
            self.ws_tax_refund['F{}'.format(row)] = '=ROUND(C{0}/(1+D{0}/100)*E{0}/100,2)'.format(row)

        self.ws_tax_refund['C{}'.format(row_number)] = '=SUM(C4:C{})'.format(row_number - 1)
        self.ws_tax_refund['F{}'.format(row_number)] = '=SUM(F4:F{})'.format(row_number - 1)
        self.ws_tax_refund['B{}'.format(row_number)].border = slash_border
        self.ws_tax_refund['E{}'.format(row_number)].border = slash_border

        # 合并需要合并单元格
        self.ws_tax_refund.merge_cells('A1:F1')

        # 打印设置
        self.ws_tax_refund.print_options.horizontalCentered = True
        self.ws_tax_refund.print_area = 'A1:F{}'.format(row_number)
        self.ws_tax_refund.page_setup.fitToWidth = 1
        self.ws_tax_refund.page_setup.orientation = "landscape"
        self.ws_tax_refund.page_margins = \
            PageMargins(left=0.25, right=0.25, top=0.75, bottom=0.75, header=0.3, footer=0.3)

    def create_summed_quotation(self):
        """创建对内总报价表"""
        self.ws_summed_quotation = self.wb.create_sheet('2.物资对内总报价表', 3)
        colum_title = ['序号', '品名', '规格', '商标', '产地', '生产年份', '{}单价{}(元人民币)'.format(
            self.project.trans, linesep), '数量', '', '{}总价{}(元人民币)'.format(self.project.trans, linesep), '备注']
        title_width = [6, 8, 50, 6, 6, 6, 14, 6, 6, 14, 10]
        colum_number = len(colum_title)
        row_number = len(self.project.commodities) + 4

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=12)
        normal_font = Font(name='宋体', size=12)
        title_font = Font(name='黑体', size=20)
        left_alignment = Alignment(
            horizontal='left',
            vertical='top',
            wrap_text=True)

        # 初始化表格
        for i in range(colum_number):
            for j in range(1, row_number):
                cell_now = self.ws_summed_quotation.cell(
                    row=j + 1, column=i + 1)
                if j < row_number - 1:
                    cell_now.border = full_border
                if j < 2:
                    cell_now.font = bold_font
                else:
                    cell_now.font = normal_font
                cell_now.alignment = ctr_alignment
        for i in range(len(title_width)):  # 修改列宽
            self.ws_summed_quotation.column_dimensions[
                self.ws_summed_quotation.cell(row=4, column=i + 1).column_letter].width = title_width[i]
        for i in range(2, row_number):  # 修改行高
            self.ws_summed_quotation.row_dimensions[
                self.ws_summed_quotation.cell(row=i, column=1).row].height = 30
        self.ws_summed_quotation.row_dimensions[row_number].height = 45

        # 创建标题行
        self.ws_summed_quotation['A1'].font = title_font
        self.ws_summed_quotation['A1'].alignment = ctr_alignment
        self.ws_summed_quotation['A1'] = '2.物资对内总报价表'
        self.ws_summed_quotation.row_dimensions[1].height = 30

        # 填写表头
        index = 0
        for i in self.ws_summed_quotation['A2':'k2'][0]:
            # print(index+1, i)
            if colum_title[index] != '':
                i.value = colum_title[index]
                # i.font = bold_font
            index += 1

        # 填入数据
        col_relate = [('A', 'A'), ('B', 'B'), ('C', 'G'), ('D', 'F'), ('E', 'L'), ('F', 'P'), ('H', 'D'),
                      ('I', 'E'), ('K', 'X')]
        for row in range(3, row_number - 1):  # 遍历行
            for col in col_relate:  # 根据对应关系设立公式
                cell_now = self.ws_summed_quotation['{}{}'.format(col[0], row)]
                cell_now.value = '=物资输入!{}{}'.format(col[1], row - 1)
                if col[0] == 'C':
                    cell_now.alignment = left_alignment
            self.ws_summed_quotation['J{}'.format(
                row)] = "='3.物资对内分项报价表'!M{}".format(row)
            self.ws_summed_quotation['J{}'.format(
                row)].number_format = '#,##0.00'
            self.ws_summed_quotation['G{}'.format(
                row)] = '=round(J{0}/I{0},2)'.format(row)
            self.ws_summed_quotation['G{}'.format(
                row)].number_format = '#,##0.00'
        self.ws_summed_quotation['A{}'.format(row_number - 1)] = '合计金额'
        self.ws_summed_quotation['C{}'.format(
            row_number - 1)] = '=SUM(J3:J{})'.format(row_number - 1)
        self.ws_summed_quotation['C{}'.format(row_number - 1)].font = bold_font
        self.ws_summed_quotation['C{}'.format(
            row_number - 1)].number_format = '¥#,##0.00'
        self.ws_summed_quotation['C{}'.format(row_number - 1)].alignment = Alignment(
            horizontal='left', vertical='center', wrap_text=True)
        self.ws_summed_quotation['A{}'.format(row_number)] = '注：'
        self.ws_summed_quotation['B{}'.format(row_number)] = "（1）本表所列{}{}单价和总价包括投标人提供上述全部物资" \
            "并承担合同规定全部义务所需的一切费用；{}（2）在备注栏中注明包装" \
            "情况,即包装的单位和数量".format(self.project.trans,
                                  self.project.destination, linesep)
        self.ws_summed_quotation['B{}'.format(row_number)].alignment = Alignment(horizontal='left',
                                                                                 vertical='center', wrap_text=True)

        # 合并需要合并单元格
        self.ws_summed_quotation.merge_cells('A1:K1')
        self.ws_summed_quotation.merge_cells('H2:I2')
        self.ws_summed_quotation.merge_cells(
            'A{0}:B{0}'.format(row_number - 1))
        self.ws_summed_quotation.merge_cells(
            'C{0}:K{0}'.format(row_number - 1))
        self.ws_summed_quotation.merge_cells('B{0}:K{0}'.format(row_number))

        # 打印设置
        self.ws_summed_quotation.print_options.horizontalCentered = True
        self.ws_summed_quotation.print_area = 'A1:K{}'.format(row_number)
        self.ws_summed_quotation.page_setup.fitToWidth = 1
        self.ws_summed_quotation.page_setup.orientation = "landscape"
        self.ws_summed_quotation.page_margins = PageMargins(left=0.25, right=0.25, top=0.75, bottom=0.75, header=0.3,
                                                            footer=0.3)

    def create_examination(self):
        """创建物资选型一览表（非法检）"""
        index = 0  # 计算表格序号
        if self.project.is_tech:
            index += 1
        if self.project.is_cc:
            index += 1
        if len(self.project.qc) > 0:
            index += 1
        self.ws_examination = self.wb.create_sheet('{}.非法检物资检验一览表'.format(index + 4), 4)
        colum_title = ['序号', '品名', 'HS编码', '数量及单位', '', '品牌', '规格型号参数', '金额', '生产厂商',
                       '供货商', '生产或供货地', '供货联系人及联系电话', '', '出厂日期', '出口港', '检验标准', '施检机构名称',
                       '', '', '备注']
        subcol_title = ['产地或供货地检验（查验）机构', '装运前核验机构', '口岸监装机构']
        title_width = [6, 14, 12, 3, 5, 8, 30, 16,
                       10, 10, 10, 10, 10, 8, 8, 15, 13, 8, 7, 6]

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        left_alignment = Alignment(
            horizontal='left',
            vertical='top',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=10)
        normal_font = Font(name='宋体', size=10)
        title_font = Font(name='黑体', size=20)

        # 合并需要合并单元格
        self.ws_examination.merge_cells('D2:E3')
        self.ws_examination.merge_cells('Q2:S2')
        self.ws_examination.merge_cells('L2:M3')
        # self.ws_examination['M2'].border = full_border
        for col in range(1, 21):
            if col not in [4, 5, 12, 13, 17, 18, 19]:
                self.ws_examination.merge_cells(
                    start_row=2, start_column=col, end_row=3, end_column=col)
        for i in range(len(title_width)):  # 修改列宽
            self.ws_examination.column_dimensions[
                self.ws_examination.cell(row=4, column=i + 1).column_letter].width = title_width[i]

        # 初始化表格
        colum_number = len(colum_title)
        row_number = len(self.project.commodities) + 3
        for i in range(colum_number):
            for j in range(row_number):
                cell_now = self.ws_examination.cell(row=j + 1, column=i + 1)
                if j > 0:  # 第一列留下给表头
                    cell_now.border = full_border
                    cell_now.font = normal_font
                    cell_now.alignment = ctr_alignment

        # 创建标题行
        self.ws_examination.merge_cells('A1:T1')
        self.ws_examination['A1'].font = title_font
        self.ws_examination['A1'].alignment = ctr_alignment
        num = ['四', '五', '六', '七']
        self.ws_examination['A1'] = '{}.非法检物资检验一览表'.format(num[index])
        self.ws_examination.row_dimensions[1].height = 30

        # 填写表头
        index = 0
        for i in self.ws_examination['A2':'T2'][0]:
            # print(index+1, i)
            if colum_title[index] != '':
                i.value = colum_title[index]
                i.font = bold_font
            index += 1
        index = 0
        for cell in self.ws_examination['Q3':'S3'][0]:
            cell.value = subcol_title[index]
            cell.font = bold_font
            index += 1

        # 填入数据
        col_relate = [('A', 'A'), ('B', 'B'), ('C', 'C'), ('D', 'D'), ('E', 'E'), ('F', 'F'), ('H', 'J'),
                      ('I', 'K'), ('J', 'L'), ('K', 'M'), ('L', 'N'), ('M', 'O'), ('N', 'P'), ('O', 'Q'), ('P', 'R'),
                      ('Q', 'S'), ('R', 'S'), ('S', 'S'), ('T', 'Y')]
        for row in range(4, row_number + 1):  # 遍历行
            for col in col_relate:  # 根据对应关系设立公式
                cell_now = self.ws_examination['{}{}'.format(col[0], row)]
                cell_now.value = '=物资输入!{}{}'.format(col[1], row - 2)
                if col[0] == 'H':
                    cell_now.number_format = '¥#,##0.00'
            self.ws_examination['G{}'.format(row)] = '="型号："&物资输入!G{0}&CHAR(10)&物资输入!H{0}'.format(row - 2)
            self.ws_examination['G{}'.format(row)].alignment = left_alignment

        # 打印设置
        self.ws_examination.print_area = 'A1:T{}'.format(row_number)
        self.ws_examination.page_setup.fitToWidth = 1
        self.ws_examination.page_setup.orientation = "landscape"
        self.ws_examination.page_margins = PageMargins(left=0.25, right=0.25, top=0.75, bottom=0.7, header=0.3,
                                                       footer=0.3)

    def create_techserve(self):
        """创建技术服务费报价表"""
        self.ws_techserve = self.wb.create_sheet('4.技术服务费报价表', 4)
        colum_title = [
            '序号',
            '费用名称',
            '美元单价',
            '人民币单价',
            '人数',
            '天/次数',
            '美元合计',
            '人民币合计']
        title_width = [6, 16, 14, 16, 8, 8, 14, 20]

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        slash_border = Border(diagonal=real_side, diagonalDown=True, left=real_side, right=real_side,
                              top=real_side, bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        right_alignment = Alignment(horizontal='right', vertical='center')
        bold_font = Font(name='宋体', bold=True, size=12)
        normal_font = Font(name='宋体', size=12)
        title_font = Font(name='黑体', size=20)
        # yellow_fill = PatternFill(
        #     fill_type='solid',
        #     start_color='FFFF00',
        #     end_color='FFFF00')

        # 初始化表格
        colum_number = len(colum_title)
        row_number = 14
        for i in range(colum_number):
            for j in range(1, row_number):  # 第一列留下给表头
                cell_now = self.ws_techserve.cell(row=j + 1, column=i + 1)
                cell_now.border = full_border
                cell_now.font = normal_font
                # if i == 6 or 7:
                #     cell_now.alignment = right_alignment
                # else:
                if i < 6:
                    cell_now.alignment = ctr_alignment
                else:
                    cell_now.alignment = right_alignment

        for i in range(len(title_width)):  # 修改列宽
            self.ws_techserve.column_dimensions[
                self.ws_techserve.cell(row=4, column=i + 1).column_letter].width = title_width[i]
        for i in range(2, row_number + 1):  # 修改行高
            self.ws_techserve.row_dimensions[
                self.ws_techserve.cell(row=i, column=1).row].height = 30
        self.ws_techserve.row_dimensions[16].height = 30

        # 打上斜线
        cell_coor = ['C3', 'C4', 'C5', 'C6', 'C7', 'D8', 'D9', 'D10', 'C11', 'D11', 'E11', 'F11', 'C12',
                     'D12', 'E12', 'F12',  'C13', 'D13', 'E13', 'F13']
        for cell in cell_coor:
            self.ws_techserve[cell].border = slash_border

        # 创建标题行
        self.ws_techserve.merge_cells('A1:H1')
        self.ws_techserve['A1'].font = title_font
        self.ws_techserve['A1'].alignment = ctr_alignment
        self.ws_techserve['A1'] = '四.技术服务费报价表'
        self.ws_techserve.row_dimensions[1].height = 40

        # 填写表头
        index = 0
        for i in self.ws_techserve['A2':'H2'][0]:
            # print(index+1, i)
            if colum_title[index] != '':
                i.value = colum_title[index]
                i.font = bold_font
            index += 1

        # 填写项目栏
        col_a = [i for i in range(1, 13)]  # 序号
        col_b = ['护照和签证手续费',
                 '防疫免疫费',
                 '技术服务人员保险费',
                 '国内交通费',
                 '国际交通费',
                 '住宿费',
                 '伙食费',
                 '津贴补贴',
                 '当地雇工费',
                 '当地设备工具材料购置或租用费',
                 '其它确需发生的费用',
                 '共计']  # 费用名称
        for index in range(3, 15):
            self.ws_techserve['A{}'.format(index)] = col_a[index - 3]
            self.ws_techserve['B{}'.format(index)] = col_b[index - 3]

        # for num in range(3):  # 填写技术服务费单价
        #     self.ws_techserve['C{}'.format(
        #         num + 6)].number_format = '$#,##0.00'
        #     self.ws_techserve['C{}'.format(
        #         num + 6)] = self.project.techinfo[2][num]

        # 格式化单元格填写数据
        for row in range(3, 15):
            self.ws_techserve['G{}'.format(row)].number_format = '$#,##0.00'
            self.ws_techserve['H{}'.format(row)].number_format = '¥#,##0.00'
            if row < 8:
                self.ws_techserve['D{}'.format(row)].number_format = '¥#,##0.00'
                self.ws_techserve['H{}'.format(row)] = '=D{0}*E{0}'.format(row)
                self.ws_techserve['G{}'.format(row)] = 0
                self.ws_techserve['F{}'.format(row)].number_format = '0'
                self.ws_techserve['F{}'.format(row)] = '-'
            if 11 > row > 7:
                self.ws_techserve['C{}'.format(row)].number_format = '$#,##0.00'
                self.ws_techserve['G{}'.format(row)] = '=C{0}*E{0}*F{0}'.format(row)
                self.ws_techserve['F{}'.format(row)].number_format = '0'
                self.ws_techserve['F{}'.format(row)] = self.project.techinfo[1]
            if 15 > row > 7:
                self.ws_techserve['H{}'.format(row)] = '=G{}*C16/100'.format(row)
            if row < 11:
                self.ws_techserve['E{}'.format(row)].number_format = '0'
                self.ws_techserve['E{}'.format(row)] = self.project.techinfo[0]

        self.ws_techserve['D3'] = 500
        self.ws_techserve['D4'] = 200
        self.ws_techserve['D5'] = 400
        self.ws_techserve['D6'] = 1000
        self.ws_techserve['D7'] = 10000
        self.ws_techserve['C8'] = 150
        self.ws_techserve['C9'] = 100
        self.ws_techserve['C10'] = 50
        self.ws_techserve['G11'] = 0
        self.ws_techserve['G12'] = 0
        self.ws_techserve['G13'] = 0
        self.ws_techserve['G14'] = '=SUM(G3:G13)'
        self.ws_techserve['H14'] = '=SUM(H3:H13)'

        # 填充备注
        self.ws_techserve['A16'] = '注：'
        self.ws_techserve['B16'] = '100美元='
        self.ws_techserve['C16'].number_format = '0.00"元人民币"'
        self.ws_techserve['C16'] = '=费用输入!F10*100'
        # self.ws_techserve['C16'].fill = yellow_fill
        self.ws_techserve['A16'].font = normal_font
        self.ws_techserve['B16'].font = normal_font
        self.ws_techserve['C16'].font = normal_font
        self.ws_techserve['A16'].alignment = right_alignment
        self.ws_techserve['B16'].alignment = right_alignment
        self.ws_techserve['C16'].alignment = Alignment(horizontal='left', vertical='center')

        # 合并需要合并单元格
        self.ws_techserve.merge_cells('C11:F11')
        self.ws_techserve.merge_cells('C12:F12')
        self.ws_techserve.merge_cells('C13:F13')
        self.ws_techserve.merge_cells('B14:F14')
        self.ws_techserve.merge_cells('C16:D16')

        # 打印设置
        self.ws_techserve.print_options.horizontalCentered = True
        self.ws_techserve.print_area = 'A1:H16'
        self.ws_techserve.page_setup.fitToWidth = 1
        # self.ws_techserve.page_setup.orientation = "landscape"
        self.ws_techserve.sheet_properties.pageSetUpPr = Quotation.fitsetup
        self.ws_techserve.page_margins = PageMargins(left=0.7, right=0.7, top=0.75, bottom=0.75, header=0.3,
                                                     footer=0.3)

    def create_lawexam(self):
        """创建物资选型一览表（法检物资）"""
        index = 0  # 计算表格序号
        if self.project.is_tech:
            index += 1
        if self.project.is_cc:
            index += 1
        self.ws_lawexam = self.wb.create_sheet('{}.法检物资检验一览表'.format(index + 4), 4)
        colum_title = ['序号', '品名', 'HS编码', '数量及单位', '', '品牌', '规格或型号', '金额', '生产厂商',
                       '供货商', '生产或供货地', '供货联系人及联系电话', '', '出厂日期', '供货地商检部门',
                       '出口港', '检验标准', '口岸监装机构', '备注']
        title_width = [6, 14, 12, 3, 5, 8, 30, 16, 10, 10, 6, 6, 10, 8, 6, 13, 8, 6]

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        left_alignment = Alignment(
            horizontal='left',
            vertical='top',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=10)
        normal_font = Font(name='宋体', size=10)
        title_font = Font(name='黑体', size=20)
        yellow_fill = PatternFill(
            fill_type='solid',
            start_color='FFFF00',
            end_color='FFFF00')

        # 合并需要合并单元格
        self.ws_lawexam.merge_cells('D2:E2')
        self.ws_lawexam.merge_cells('L2:M2')
        # self.ws_examination['M2'].border = full_border
        for i in range(len(title_width)):  # 修改列宽
            self.ws_lawexam.column_dimensions[
                self.ws_lawexam.cell(row=4, column=i + 1).column_letter].width = title_width[i]

        # 初始化表格
        colum_number = len(colum_title)
        row_number = len(self.project.qc) + 2
        for i in range(colum_number):
            for j in range(row_number):
                cell_now = self.ws_lawexam.cell(row=j + 1, column=i + 1)
                if j > 0:  # 第一列留下给表头
                    cell_now.border = full_border
                    cell_now.font = normal_font
                    cell_now.alignment = ctr_alignment

        # 创建标题行
        self.ws_lawexam.merge_cells('A1:S1')
        self.ws_lawexam['A1'].font = title_font
        self.ws_lawexam['A1'].alignment = ctr_alignment
        num = ['三', '四', '五', '六', '七']
        self.ws_lawexam['A1'] = '{}.法检物资检验一览表'.format(num[index])
        self.ws_lawexam.row_dimensions[1].height = 30

        # 填写表头
        index = 0
        for i in self.ws_lawexam['A2':'S2'][0]:
            # print(index+1, i)
            if colum_title[index] != '':
                i.value = colum_title[index]
                i.font = bold_font
            index += 1

        # 填入数据
        col_relate = [('A', 'A'), ('B', 'B'), ('C', 'C'), ('D', 'D'), ('E', 'E'), ('F', 'F'), ('H', 'J'),
                      ('I', 'K'), ('J', 'L'), ('K', 'M'), ('L', 'N'), ('M', 'O'), ('N', 'P'), ('P', 'Q'),
                      ('Q', 'R'), ('R', 'S'), ('S', 'Y')]
        index = 0
        for row in range(3, row_number + 1):  # 遍历行
            for col in col_relate:  # 根据对应关系设立公式
                cell_now = self.ws_lawexam['{}{}'.format(col[0], row)]
                cell_now.value = '=物资输入!{}{}'.format(
                    col[1], self.project.qc[index] + 1)
                if col[0] == 'H':
                    cell_now.number_format = '¥#,##0.00'
            self.ws_lawexam['G{}'.format(row)] = '="型号："&物资输入!G{0}&CHAR(10)&物资输入!H{0}'.\
                format(self.project.qc[index] + 1)
            self.ws_lawexam['G{}'.format(row)].alignment = left_alignment
            self.ws_lawexam['O{}'.format(row)] = '=K{}&"海关"'.format(row)
            self.ws_lawexam['O{}'.format(row)].fill = yellow_fill
            index += 1
        num = 0
        for row in self.project.qc:  # 在非法检物资中删除
            self.ws_examination.delete_rows(row - num + 3)
            num += 1

        # 打印设置
        self.ws_lawexam.print_options.horizontalCentered = True
        self.ws_lawexam.print_area = 'A1:S{}'.format(row_number)
        self.ws_lawexam.page_setup.fitToWidth = 1
        self.ws_lawexam.page_setup.orientation = "landscape"
        self.ws_lawexam.page_margins = PageMargins(left=0.25, right=0.25, top=0.75, bottom=0.75, header=0.3,
                                                   footer=0.3)

    def create_training(self):
        """创建来华培训费报价表"""
        index = 0
        if self.project.is_tech:
            index += 1
        self.ws_training = self.wb.create_sheet('{}.来华培训费报价表'.format(index + 4), 4)
        colum_title = [
            '序号',
            '费用名称',
            '',
            '费用计算方式',
            '',
            '',
            '人民币（元）',
            '其中含购汇人民币限额']
        title_width = [6, 15, 8, 14, 7, 12, 14, 12]
        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        slash_border = Border(diagonal=real_side, diagonalDown=True, left=real_side, right=real_side,
                              top=real_side, bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrapText=True)
        right_alignment = Alignment(
            horizontal='right',
            vertical='center',
            wrapText=True)
        left_alignment = Alignment(
            horizontal='left',
            vertical='center',
            wrapText=True)
        bold_font = Font(name='宋体', bold=True, size=12)
        normal_font = Font(name='宋体', size=12)
        title_font = Font(name='黑体', bold=True, size=14)
        yellow_fill = PatternFill(
            fill_type='solid',
            start_color='FFFF00',
            end_color='FFFF00')

        # 初始化表格
        colum_number = len(colum_title)
        row_number = 20
        for i in range(colum_number):
            for j in range(1, row_number):  # 第一列留下给表头
                cell_now = self.ws_training.cell(row=j + 1, column=i + 1)
                if j < 17:  # 给主体cell设置样式
                    cell_now.border = full_border
                    cell_now.font = normal_font
                    if i == 6:  # 数字列右对齐
                        cell_now.alignment = right_alignment
                    else:
                        cell_now.alignment = ctr_alignment
                else:  # 最后两行左对齐
                    cell_now.font = normal_font
                    cell_now.alignment = left_alignment
        for i in range(len(title_width)):  # 修改列宽
            self.ws_training.column_dimensions[
                self.ws_training.cell(row=4, column=i + 1).column_letter].width = title_width[i]
        for i in range(2, row_number):  # 修改行高
            self.ws_training.row_dimensions[
                self.ws_training.cell(row=i, column=1).row].height = 20
        self.ws_training.row_dimensions[20].height = 30

        # 打上斜线
        cell_coor = ['D14', 'E14', 'F14']
        for cell in cell_coor:
            self.ws_training[cell].border = slash_border

        # 创建标题行
        self.ws_training['A1'].font = title_font
        self.ws_training['A1'].alignment = ctr_alignment
        num = ['四', '五', '六', '七']
        self.ws_training['A1'] = '{}.来华培训费报价表'.format(num[index])
        self.ws_training.row_dimensions[1].height = 30
        self.ws_training.merge_cells('A1:H1')

        # 填写表头
        index = 0
        for i in self.ws_training['A2':'H2'][0]:
            # print(index, i)
            if colum_title[index] != '':
                i.value = colum_title[index]
                i.font = bold_font
            index += 1
        self.ws_training['D3'].value = '标准'
        self.ws_training['E3'].value = '人数'
        self.ws_training['F3'].value = '天（次）数'
        self.ws_training['D3'].font = bold_font
        self.ws_training['E3'].font = bold_font
        self.ws_training['F3'].font = bold_font

        # 填写数据
        col_a = ['一', '二', 1, 2, 3, 4, 5, 6, '三', '四', 1, 2, '', '五', '']  # 序号
        col_b = ['培训费', '接待费', '日常伙食费', '住宿费', '宴请费', '零用费', '小礼品费', '人身意外伤害保险',
                 '国际旅费', '管理费', '承办管理费', '管理人员费', '', '合计']  # 费用名称
        for index in range(4, 18):  # 填写前两列
            self.ws_training['A{}'.format(index)] = col_a[index - 4]
            if isinstance(col_a[index - 4], int):
                self.ws_training['A{}'.format(
                    index)].alignment = right_alignment
            self.ws_training['B{}'.format(index)] = col_b[index - 4]
        self.ws_training['C15'].value = '伙食费'
        self.ws_training['C16'].value = '住宿费'

        # 填写E列
        for num in range(6, 13):  # 填写培训人数
            self.ws_training['E{}'.format(num)].number_format = '0'
            self.ws_training['E{}'.format(
                num)].value = self.project.training_num
        self.ws_training['E4'].number_format = '0'
        self.ws_training['E4'].value = self.project.training_num
        self.ws_training['E15'].number_format = '0'
        num = self.project.training_num
        if num < 10:
            res = 1
        else:
            if num % 10 == 0:
                res = num / 10
            else:
                res = num / 10 + 1
        self.ws_training['E15'].value = res
        self.ws_training['E16'].number_format = '0'
        self.ws_training['E16'].value = '=E15'

        # 填写D列
        for num in [4, 6, 7, 9, 15, 16]:
            self.ws_training['D{}'.format(num)].number_format = '0"元/人*天"'
        for num in range(10, 13):
            self.ws_training['D{}'.format(num)].number_format = '0"元/人"'
        self.ws_training['D8'].number_format = '0"元/人*次"'
        self.ws_training['D4'] = 320
        self.ws_training['D6'] = 140
        self.ws_training['D7'] = 300
        self.ws_training['D8'] = 150
        self.ws_training['D9'] = 80
        self.ws_training['D10'] = 200
        self.ws_training['D11'] = 100
        self.ws_training['D12'] = 0
        self.ws_training['D12'].fill = yellow_fill
        self.ws_training['D15'] = 140
        self.ws_training['D16'] = 300

        # 填写F列
        for i in [4, 6, 7, 8, 9, 10, 11, 12, 15, 16]:
            self.ws_training['F{}'.format(i)].number_format = '0'
        self.ws_training['F4'] = self.project.training_days
        self.ws_training['F6'] = self.project.training_days
        self.ws_training['F7'] = self.project.training_days - 1
        self.ws_training['F9'] = self.project.training_days
        self.ws_training['F11'] = '-'
        self.ws_training['F12'] = '-'
        self.ws_training['F15'] = self.project.training_days
        self.ws_training['F16'] = self.project.training_days
        if self.project.is_lowprice:
            self.ws_training['F8'] = 0
            self.ws_training['F10'] = 0
        else:
            self.ws_training['F8'] = 1
            self.ws_training['F10'] = 1

        # 填写G列
        for i in range(4, 18):
            self.ws_training['G{}'.format(i)].number_format = '¥#,##0.00'
            if i in [5, 13]:
                pass
            elif i in (11, 12):
                self.ws_training['G{}'.format(i)] = '=D{0}*E{0}'.format(i)
            elif i == 14:
                self.ws_training['G{}'.format(
                    i)] = '=ROUND((SUM(G4,G6:G11))*0.06,2)'
            elif i == 17:
                self.ws_training['G{}'.format(i)] = '=SUM(G4:G16)'
            else:
                self.ws_training['G{}'.format(i)] = '=D{0}*E{0}*F{0}'.format(i)

        # 填充备注
        self.ws_training['A19'] = '注：'
        self.ws_training['B19'] = '（1）100美元='
        self.ws_training['C19'].number_format = '0.00"元人民币"'
        self.ws_training['C19'] = '=费用输入!F10'
        # self.ws_training['C19'].fill = yellow_fill
        self.ws_training['B20'] = '（2）上述费用参照财政部（2008）第2号文举办援外培训班费用开支标准和财务管理办法给定的费用标准报价'

        # 合并需要合并单元格
        self.ws_training.merge_cells('D2:F2')
        self.ws_training.merge_cells('A2:A3')
        self.ws_training.merge_cells('B2:C3')
        self.ws_training.merge_cells('G2:G3')
        self.ws_training.merge_cells('H2:H3')
        self.ws_training.merge_cells('D5:G5')
        self.ws_training.merge_cells('D13:G13')
        self.ws_training.merge_cells('D14:F14')
        self.ws_training.merge_cells('B17:F17')
        self.ws_training.merge_cells('C19:D19')
        self.ws_training.merge_cells('B20:H20')
        for i in range(4, 15):
            self.ws_training.merge_cells('B{0}:C{0}'.format(i))
        self.ws_training.merge_cells('B15:B16')
        self.ws_training.merge_cells('A15:A16')

        # 打印设置
        self.ws_training.print_options.horizontalCentered = True
        self.ws_training.print_area = 'A1:H{}'.format(row_number)
        self.ws_training.page_setup.fitToWidth = 1
        # self.ws_training.page_setup.orientation = "landscape"
        self.ws_training.sheet_properties.pageSetUpPr = Quotation.fitsetup
        self.ws_training.page_margins = PageMargins(left=0.25, right=0.25, top=0.75, bottom=0.75, header=0.3,
                                                    footer=0.3)

    def create_isolist(self):
        """创建三体系一览表"""
        index = 0
        if self.project.is_tech:
            index += 1
        if self.project.is_qa:
            index += 1
        self.ws_isolist = self.wb.create_sheet('{}.三体系一览表'.format(index + 8), -1)
        colum_title = ['序号', '物资名称', '生产企业名称', '招标要求', '投标响应', '认证文件编号']
        title_width = [5, 15, 25, 35, 9, 55]
        colum_number = len(colum_title)
        row_number = len(self.project.commodities) + 2

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=12)
        normal_font = Font(name='宋体', size=10)
        title_font = Font(name='黑体', bold=True, size=18)

        # 初始化表格
        for i in range(colum_number):
            for j in range(row_number):
                cell_now = self.ws_isolist.cell(row=j + 1, column=i + 1)
                if j > 0:
                    cell_now.border = full_border
                if j > 1:
                    cell_now.font = normal_font
                if i == 5 and j > 1:
                    cell_now.alignment = Alignment(
                        horizontal='left', vertical='center', wrap_text=True)
                else:
                    cell_now.alignment = ctr_alignment
        for i in range(len(title_width)):  # 修改列宽
            self.ws_isolist.column_dimensions[
                self.ws_isolist.cell(row=4, column=i + 1).column_letter].width = title_width[i]
        self.ws_isolist.row_dimensions[2].height = 40

        # 创建标题行
        num = ['八', '九', '十', '十一', '十二', '十三']  # 存放中文序号
        self.ws_isolist.merge_cells('A1:F1')
        self.ws_isolist.merge_cells('D3:D{}'.format(row_number))
        self.ws_isolist['A1'].font = title_font
        self.ws_isolist['A1'].alignment = ctr_alignment
        self.ws_isolist['A1'] = '{}.物资生产企业质量管理、环境管理和职业健康安全管理体系认证一览表'.format(num[index])
        self.ws_isolist.row_dimensions[1].height = 50

        # 填写表头
        index = 0
        for i in self.ws_isolist['A2':'F2'][0]:
            # print(index+1, i)
            if colum_title[index] != '':
                i.value = colum_title[index]
                i.font = bold_font
            index += 1

        # 填入数据
        col_relate = [('A', 'A'), ('B', 'B'), ('C', 'K'), ('F', 'V')]
        for row in range(3, row_number + 1):  # 遍历行
            for col in col_relate:  # 根据对应关系设立公式
                cell_now = self.ws_isolist['{}{}'.format(col[0], row)]
                cell_now.value = '=物资输入!{}{}'.format(col[1], row - 1)
            self.ws_isolist['E{}'.format(row)] = '响应'
        self.ws_isolist['D3'] = '1. 在满足清单参数要求的前提下，鼓励本项目投标人各项物资均选用具备质量管理体系、环境管理体系和' \
                                '职业健康与安全管理体系认证的企业生产的物资。{}2. 需提交有效的管理体系认证证明文件为：管理体系认证证书复印件。'.format(linesep)

        # 打印设置
        self.ws_isolist.print_options.horizontalCentered = True
        self.ws_isolist.print_area = 'A1:F{}'.format(row_number)
        self.ws_isolist.page_setup.fitToWidth = 1
        self.ws_isolist.sheet_properties.pageSetUpPr = Quotation.fitsetup
        self.ws_isolist.page_margins = Quotation.margin

    def create_conservlist(self):
        """创建节能认证一览表"""
        index = 1
        if self.project.is_tech:
            index += 1
        if self.project.is_qa:
            index += 1
        self.ws_conservlist = self.wb.create_sheet('{}.节能产品一览表'.format(index + 8), -1)
        colum_title = ['序号', '物资名称', '品牌和型号', '招标要求', '投标响应', '认证文件编号']
        title_width = [5, 15, 25, 35, 9, 20]
        colum_number = len(colum_title)
        row_number = len(self.project.commodities) + 2

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=12)
        normal_font = Font(name='宋体', size=10)
        title_font = Font(name='黑体', bold=True, size=18)

        # 初始化表格
        for i in range(colum_number):
            for j in range(row_number):
                cell_now = self.ws_conservlist.cell(row=j + 1, column=i + 1)
                if j > 0:
                    cell_now.border = full_border
                if j > 1:
                    cell_now.font = normal_font
                if i == 5 and j > 1:
                    cell_now.alignment = Alignment(
                        horizontal='left', vertical='center', wrap_text=True)
                else:
                    cell_now.alignment = ctr_alignment
        for i in range(len(title_width)):  # 修改列宽
            self.ws_conservlist.column_dimensions[
                self.ws_conservlist.cell(row=4, column=i + 1).column_letter].width = title_width[i]
        self.ws_conservlist.row_dimensions[2].height = 40

        # 创建标题行
        num = ['八', '九', '十', '十一', '十二', '十三']  # 存放中文序号
        self.ws_conservlist.merge_cells('A1:F1')
        self.ws_conservlist.merge_cells('D3:D{}'.format(row_number))
        self.ws_conservlist['A1'].font = title_font
        self.ws_conservlist['A1'].alignment = ctr_alignment
        self.ws_conservlist['A1'] = '{}.节能产品一览表'.format(num[index])
        self.ws_conservlist.row_dimensions[1].height = 50

        # 填写表头
        index = 0
        for i in self.ws_conservlist['A2':'F2'][0]:
            # print(index+1, i)
            if colum_title[index] != '':
                i.value = colum_title[index]
                i.font = bold_font
            index += 1

        # 填入数据
        col_relate = [('A', 'A'), ('B', 'B'), ('C', 'D'), ('F', 'W')]
        for row in range(3, row_number + 1):  # 遍历行
            for col in col_relate:  # 根据对应关系设立公式
                cell_now = self.ws_conservlist['{}{}'.format(col[0], row)]
                if col[0] == 'C':
                    cell_now.value = "='0.物资选型一览表'!D{}".format(row)
                else:
                    cell_now.value = '=物资输入!{}{}'.format(col[1], row - 1)
            self.ws_conservlist['E{}'.format(row)] = '响应'
        self.ws_conservlist['D3'] = '1. 在满足清单参数要求的前提下，鼓励本项目投标人各项物资均选用具备节能产品认证的物资。{}' \
             '2. 需提交有效的节能产品认证证明文件为：节能产品认证证书复印件（提交的证书须符合《市场监管总局关于发布参与实施政府' \
                                    '采购节能产品、环境标志产品认证机构名录的公告》（2019年第16号）等文件要求）。'.format(linesep)

        # 打印设置
        self.ws_conservlist.print_options.horizontalCentered = True
        self.ws_conservlist.print_area = 'A1:F{}'.format(row_number)
        self.ws_conservlist.page_setup.fitToWidth = 1
        self.ws_conservlist.sheet_properties.pageSetUpPr = Quotation.fitsetup
        self.ws_conservlist.page_margins = Quotation.margin

    def create_eplist(self):
        """创建环保认证一览表"""
        index = 2
        if self.project.is_tech:
            index += 1
        if self.project.is_qa:
            index += 1
        self.ws_eplist = self.wb.create_sheet('{}.环境标志产品一览表'.format(index + 8), -1)
        colum_title = ['序号', '物资名称', '品牌和型号', '招标要求', '投标响应', '认证文件编号']
        title_width = [5, 15, 25, 35, 9, 20]
        colum_number = len(colum_title)
        row_number = len(self.project.commodities) + 2

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=12)
        normal_font = Font(name='宋体', size=10)
        title_font = Font(name='黑体', bold=True, size=18)

        # 初始化表格
        for i in range(colum_number):
            for j in range(row_number):
                cell_now = self.ws_eplist.cell(row=j + 1, column=i + 1)
                if j > 0:
                    cell_now.border = full_border
                if j > 1:
                    cell_now.font = normal_font
                if i == 5 and j > 1:
                    cell_now.alignment = Alignment(
                        horizontal='left', vertical='center', wrap_text=True)
                else:
                    cell_now.alignment = ctr_alignment
        for i in range(len(title_width)):  # 修改列宽
            self.ws_eplist.column_dimensions[
                self.ws_eplist.cell(row=4, column=i + 1).column_letter].width = title_width[i]
        self.ws_eplist.row_dimensions[2].height = 40

        # 创建标题行
        num = ['八', '九', '十', '十一', '十二', '十三']  # 存放中文序号
        self.ws_eplist.merge_cells('A1:F1')
        self.ws_eplist.merge_cells('D3:D{}'.format(row_number))
        self.ws_eplist['A1'].font = title_font
        self.ws_eplist['A1'].alignment = ctr_alignment
        self.ws_eplist['A1'] = '{}.环境标志产品一览表'.format(num[index])
        self.ws_eplist.row_dimensions[1].height = 50

        # 填写表头
        index = 0
        for i in self.ws_eplist['A2':'F2'][0]:
            # print(index+1, i)
            if colum_title[index] != '':
                i.value = colum_title[index]
                i.font = bold_font
            index += 1

        # 填入数据
        col_relate = [('A', 'A'), ('B', 'B'), ('C', 'D'), ('F', 'X')]
        for row in range(3, row_number + 1):  # 遍历行
            for col in col_relate:  # 根据对应关系设立公式
                cell_now = self.ws_eplist['{}{}'.format(col[0], row)]
                if col[0] == 'C':
                    cell_now.value = "='0.物资选型一览表'!D{}".format(row)
                else:
                    cell_now.value = '=物资输入!{}{}'.format(col[1], row - 1)
            self.ws_eplist['E{}'.format(row)] = '响应'
        re = '1. 在满足清单参数要求的前提下，鼓励本项目投标人各项物资均选用具备环境标志产品认证的物资。{}2.' \
             '需提交有效的环境标志产品认证证明文件为：环境标志产品认证证书复印件（提交的证书须符合《市场监管总局关于发布参与' \
             '实施政府采购节能产品、环境标志产品认证机构名录的公告》（2019年第16号）等文件要求）。'.format(linesep)
        self.ws_eplist['D3'] = re

        # 打印设置
        self.ws_eplist.print_options.horizontalCentered = True
        self.ws_eplist.print_area = 'A1:F{}'.format(row_number)
        self.ws_eplist.page_setup.fitToWidth = 1
        self.ws_eplist.sheet_properties.pageSetUpPr = Quotation.fitsetup
        self.ws_eplist.page_margins = Quotation.margin

    def create_lob(self):
        '''创建投标函'''
        self.ws_lob = self.wb.create_sheet('1.投标函', 3)
        colum_title = ['序号', '费用项目', '金额{}（小写人民币元）'.format(linesep), '备注']
        title_width = [8, 35, 25, 50]
        row_hight = [50, 30, 68, 40, 65, 80, 40, 40, 30, 44, 200, 90, 20, 60, 100]
        colum_number = len(colum_title)
        row_number = 15
        if self.project.is_tech:
            row_number += 1
            row_hight.insert(5, 40)
        if self.project.is_cc:
            row_number += 1
            row_hight.insert(5, 40)

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        right_alignment = Alignment(
            horizontal='right',
            vertical='center',
            wrap_text=True)
        left_alignment = Alignment(
            horizontal='left',
            vertical='center',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=14)
        normal_font = Font(name='宋体', size=14)
        title_font = Font(name='宋体', bold=True, size=20)

        # 初始化表格
        for i in range(1, colum_number + 1):
            for j in range(1, row_number + 1):
                cell_now = self.ws_lob.cell(row=j, column=i)
                if row_number - 6 > j > 3:
                    cell_now.border = full_border
                    cell_now.font = normal_font
                    if i == 3:
                        cell_now.alignment = right_alignment
                        cell_now.number_format = '¥#,##0.00'
                    elif i == 1:
                        cell_now.alignment = ctr_alignment
                    else:
                        cell_now.alignment = left_alignment
                else:
                    cell_now.font = normal_font
                    cell_now.alignment = left_alignment
                if j == 4:
                    cell_now.font = bold_font
                    cell_now.alignment = ctr_alignment
                if j == row_number - 3:
                    cell_now.font = bold_font
                if j == row_number -6:
                    cell_now.alignment = ctr_alignment

        for i in range(len(title_width)):  # 修改列宽
            self.ws_lob.column_dimensions[
                self.ws_lob.cell(row=4, column=i + 1).column_letter].width = title_width[i]
        for row in range(len(row_hight)):  # 修改行高
            self.ws_lob.row_dimensions[row + 1].height = row_hight[row]

        # 创建标题行
        self.ws_lob['A1'].font = title_font
        self.ws_lob['A1'].alignment = ctr_alignment
        self.ws_lob['A1'] = '一.投标函'

        # 填写表头
        index = 0
        for i in self.ws_lob['A4':'D4'][0]:
            # print(index+1, i)
            if colum_title[index] != '':
                i.value = colum_title[index]
            index += 1

        # 填写数据
        self.ws_lob['A5'] = '一'
        self.ws_lob['B5'] = "全部物资{}{}".format(self.project.trans, self.project.destination)
        self.ws_lob['C5'] = "='1.投标报价总表'!C4"
        self.ws_lob['D5'] = \
            "含商品购买价款、国内运杂费、包装费、保管费、物资检验费、运输保险费、国外运费、资金占用成本、合理利润、税金"

        if self.project.is_tech:
            self.ws_lob['C6'] = "='4.技术服务费报价表'!H14"
            self.ws_lob['B6'] = '技术服务费'
            self.ws_lob['D6'] = '="含："&TEXT(\'4.技术服务费报价表\'!G14,"#,##0.00")&' \
                                '"美元"&CHAR(10)&"汇率：100美元="&\'4.技术服务费报价表\'!C16&"元人民币"'
            if self.project.is_cc:
                self.ws_lob['C7'] = "='5.来华培训费报价表'!G17"
                self.ws_lob['B7'] = '来华培训费'
        elif self.project.is_cc:
            self.ws_lob['C6'] = "='4.来华培训费报价表'!G17"
            self.ws_lob['B6'] = '来华培训费'

        no_seq = ['二', '三', '四', '五']
        for i in range(6, row_number - 7):
            self.ws_lob['A{}'.format(i)] = no_seq[i - 6]
        self.ws_lob["B{}".format(row_number - 9)] = "其他费用"
        self.ws_lob["C{}".format(row_number - 9)] = "=费用输入!J17"
        self.ws_lob['D{}'.format(row_number - 9)] = '="含：管理费用"&TEXT(费用输入!G15,"#,##0.00"&"元"&CHAR(10)&' \
                                                    '"风险预涨费费用"&TEXT(费用输入!J15,"#,##0.00")&"元"&CHAR(10)&' \
                                                    '"防恐措施费"&TEXT(费用输入!J16,"#,##0.00")&"元"&CHAR(10)&' \
                                                    '"大型机电设备跟踪反馈工作费"&TEXT(费用输入!G16,"#,##0.00")&"元")'
        self.ws_lob['B{}'.format(row_number - 8)] = '《供货清单（一）》中各项物资增值税退抵税额'
        self.ws_lob['C{}'.format(row_number - 8)] = \
            "='3.各项物资增值税退抵税额表'!F{}".format(len(self.project.commodities) + 4)
        self.ws_lob['B{}'.format(row_number - 7)] = '合计金额'
        self.ws_lob['C{}'.format(row_number - 7)] = "=SUM(C5:C{})-C{}".format(
            row_number - 9, row_number - 8)
        self.ws_lob['A2'] = '中国国际经济技术交流中心:'
        self.ws_lob['A3'] = '    一、我公司已仔细研究了{}的招标文件（标书编号：{}）的全部内容，愿意以下表所列金额承担本项目全部' \
                            '实施任务和内部总承包合同规定的各项义务：'.format(self.project.name, self.project.code)
        if row_number == 17:
            self.ws_lob['A{}'.format(row_number - 6)] = '（注：合计金额=一+二+三+四-五）'
        elif row_number == 16:
            self.ws_lob['A{}'.format(row_number - 6)] = '（注：合计金额=一+二+三-四）'
        else:
            self.ws_lob['A{}'.format(row_number - 6)] = '（注：合计金额=一+二-三）'
        self.ws_lob['A{}'.format(row_number - 5)] = \
            '    二、如果我公司中标，我公司保证于{}将全部物资发运完毕。'.format(self.project.trans_time)
        self.ws_lob['A{}'.format(row_number - 4)] = \
            '''    三、如果我公司中标，我公司将提交金额为中标金额10%的银行保函作为履约保证金，或履约和无缺陷质量保证金。
    四、我公司同意自你中心收到本投标书之日起的180天内，本投标书及我公司作出的补充澄清将始终对我公司具有约束力。如我公司中标，至我公司完成本项目内部实施合同规定由我公司履行的全部义务止，本投标书及我公司作出的补充澄清将始终对我公司具有约束力。
    五、我公司一旦收到中标通知书，将在30天内向你中心提交履约保证金银行保函/履约和无缺陷质量保证金银行保函，并派出法定代表人或其授权代表到你中心签署内部总承包合同。如果逾期不提交上述保函或不与你中心签约，即自动放弃中标资格。你中心有权重新授标。
    六、我公司理解，如我公司未中标，你方有权不作任何解释。
    七、我公司承诺，不以任何形式干扰评标工作。'''
        self.ws_lob['A{}'.format(row_number - 3)] = '    八、我公司已对本投标文件全部内容（包括证明物资及其生产供货企业' \
                                                    '以及我公司各项服务任务符合招标文件要求的技术支持资料）进行核实，' \
                                                    '保证全部内容均真实有效，并承诺按照采购代理机构的要求在接到质询通知后' \
                                                    '3个工作日内提供相关文件资料的正本备查核验；' \
                                                    '如无法按时提供相关材料正本，无条件接受丧失中标资格的后果' \
                                                    '以及你中心根据相关法律法规、规章制度和本项目招标文件作出的一切处理决定。'
        self.ws_lob['C{}'.format(row_number - 2)] = '公司名称（盖公章）：中国海外经济合作有限公司'
        self.ws_lob['C{}'.format(row_number - 1)] = '法人代表或其授权代表：'
        self.ws_lob['C{}'.format(row_number)] = '''地  址：北京市西城区阜外大街6号
电  话：010-68013962
传  真：010-68059153
项目负责人：张帅
{}'''.format(self.project.date)

        # 合并单元格
        self.ws_lob.merge_cells('A2:D2')
        self.ws_lob.merge_cells('A1:D1')
        self.ws_lob.merge_cells('A3:D3')
        for row in range(row_number - 3, row_number - 7, -1):
            self.ws_lob.merge_cells('A{0}:D{0}'.format(row))
        for row in range(row_number, row_number - 3, -1):
            self.ws_lob.merge_cells('C{0}:D{0}'.format(row))

        # 打印设置
        self.ws_lob.print_options.horizontalCentered = True
        self.ws_lob.print_area = 'A1:D{}'.format(row_number)
        self.ws_lob.page_setup.fitToWidth = 1
        self.ws_lob.sheet_properties.pageSetUpPr = Quotation.fitsetup
        self.ws_lob.page_margins = Quotation.margin

    def create_self_exam(self):
        """创建自检验收表格"""
        self.ws_self_exam = self.wb.create_sheet('自检验收表', -1)
        colum_title = ['序号', '品名', '供货商', '生产或供货地', '拟进行自检验收时间', '人数']
        title_width = [5, 14, 18, 10, 14, 6]

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=9)
        normal_font = Font(name='宋体', size=9)

        for i in range(len(title_width)):  # 修改列宽
            self.ws_self_exam.column_dimensions[
                self.ws_self_exam.cell(row=4, column=i + 1).column_letter].width = title_width[i]

        # 初始化表格
        colum_number = len(colum_title)
        row_number = len(self.project.commodities) + 1
        for i in range(colum_number):
            for j in range(row_number):
                cell_now = self.ws_self_exam.cell(row=j + 1, column=i + 1)
                cell_now.border = full_border
                cell_now.font = normal_font
                cell_now.alignment = ctr_alignment

        self.ws_self_exam.row_dimensions[1].height = 30

        # 填写表头
        index = 0
        for i in self.ws_self_exam['A1':'F1'][0]:
            # print(index+1, i)
            i.value = colum_title[index]
            i.font = bold_font
            index += 1

        # 填入数据
        col_relate = [('A', 'A'), ('B', 'B'), ('C', 'L'), ('D', 'M')]
        for row in range(2, row_number + 1):  # 遍历行
            for col in col_relate:  # 根据对应关系设立公式
                cell_now = self.ws_self_exam['{}{}'.format(col[0], row)]
                cell_now.value = '=物资输入!{}{}'.format(col[1], row)
            self.ws_self_exam['E{}'.format(row)] = '物资交付当天{}预计用时2天'.format(linesep)
            self.ws_self_exam['F{}'.format(row)] = '2人'

    def create_3rd_exam(self):
        """创建第三方检验表格"""
        self.ws_self_exam = self.wb.create_sheet('第三方检验表', -1)
        colum_title = ['序号', '品名', '厂家交货期', '自检验收时间', '第三方产地检验时间',
                       '装运前检验时间', '口岸监装时间', '发运时间']
        title_width = [5, 14, 14, 16, 16, 18, 18, 10]

        days = ''
        for d in self.project.trans_time:
            if d.isdigit():
                days += d
        days = days + '天内'

        # 设置基本的样式
        real_side = Side(style='thin')
        full_border = Border(
            left=real_side,
            right=real_side,
            top=real_side,
            bottom=real_side)
        ctr_alignment = Alignment(
            horizontal='center',
            vertical='center',
            wrap_text=True)
        bold_font = Font(name='宋体', bold=True, size=9)
        normal_font = Font(name='宋体', size=9)

        for i in range(len(title_width)):  # 修改列宽
            self.ws_self_exam.column_dimensions[
                self.ws_self_exam.cell(row=4, column=i + 1).column_letter].width = title_width[i]

        # 初始化表格
        colum_number = len(colum_title)
        row_number = len(self.project.commodities) + 1
        for i in range(colum_number):
            for j in range(row_number):
                cell_now = self.ws_self_exam.cell(row=j + 1, column=i + 1)
                cell_now.border = full_border
                cell_now.font = normal_font
                cell_now.alignment = ctr_alignment

        # 填写表头
        index = 0
        for i in self.ws_self_exam['A1':'H1'][0]:
            i.value = colum_title[index]
            i.font = bold_font
            index += 1

        # 填入数据
        col_relate = [('A', 'A'), ('B', 'B'), ('C', 'T')]
        for row in range(2, row_number + 1):  # 遍历行
            for col in col_relate:  # 根据对应关系设立公式
                cell_now = self.ws_self_exam['{}{}'.format(col[0], row)]
                cell_now.value = '=物资输入!{}{}'.format(col[1], row)
            self.ws_self_exam['D{}'.format(row)] = '物资交付后2天完成'
            self.ws_self_exam['E{}'.format(row)] = '物资出厂前5天完成'
            self.ws_self_exam['F{}'.format(row)] = '口岸发运前3-5天完成'
            self.ws_self_exam['G{}'.format(row)] = '口岸发运前3-5天完成'
            self.ws_self_exam['H{}'.format(row)] = days


class Content(object):
    """通过project实例创建目录"""

    # 设置公用样式
    title_font = Font(name='宋体', size=24, bold=True)
    header_font = Font(name='仿宋_GB2312', size=14, bold=True)
    normal_font = Font(name='仿宋_GB2312', size=14)
    header_border = Border(bottom=Side(style='medium'))
    normal_border = Border(bottom=Side(style='thin', color='80969696'))
    ctr_alignment = Alignment(
        horizontal='center',
        vertical='center',
        wrap_text=True)
    left_alignment = Alignment(
        horizontal='left',
        vertical='center',
        wrap_text=True,
        indent=1)
    margin = PageMargins()

    def __init__(self, project):
        self.project = project
        self.wb = Workbook()
        self.ws_lob = None
        self.ws_tech = None
        self.ws_qual = None
        self.ws_eco = None
        self.ws_com = None

    def create_all(self):
        """生成目录总方法"""
        self.create_qual()
        self.create_com()
        self.create_eco()
        self.create_tech()
        self.create_lob()
        self.wb.save('目录—{}.xlsx'.format(self.project.name))

    def create_lob(self):
        """创建投标函目录"""
        self.ws_lob = self.wb.create_sheet('投标函', 0)
        col_titles = ['序号', '内容', '页码']
        content = [['一', '投标函'], ['二', '法定代表人身份证明书'], ['三', '法定代表人授权书'],
                   ['四', '守法廉政承诺书'], ['五', '企业内控承诺'], ['六', '投标保证金银行保函']]
        col_width = [10, 60, 10]
        col_num = 3
        row_num = 8

        # 初始化表格
        for i in range(row_num):
            for j in range(col_num):
                cell_now = self.ws_lob.cell(row=i + 1, column=j + 1)
                self.ws_lob.row_dimensions[i + 1].height = 45  # 修改行高
                if i > 0:
                    if i == 1:
                        cell_now.font = Content.header_font
                        cell_now.alignment = Content.ctr_alignment
                        cell_now.border = Content.header_border
                        cell_now.value = col_titles[j]
                    else:
                        cell_now.font = Content.normal_font
                        if j == 1:
                            cell_now.alignment = Content.left_alignment
                            cell_now.value = content[i - 2][1]
                        else:
                            cell_now.alignment = Content.ctr_alignment
                            if j == 0:
                                cell_now.value = content[i - 2][0]
                            elif j == 2:
                                cell_now.value = i - 1
                        if i != row_num - 1:
                            cell_now.border = Content.normal_border
        letters = string.ascii_uppercase
        for i in range(col_num):  # 修改列宽
            self.ws_lob.column_dimensions[letters[i]].width = col_width[i]

        # 填写抬头
        self.ws_lob.merge_cells('A1:C1')
        header = self.ws_lob['A1']
        header.font = Content.title_font
        header.alignment = Content.ctr_alignment
        header.value = '目  录'
        self.ws_lob.row_dimensions[1].height = 50

        # 打印设置
        self.ws_lob.print_options.horizontalCentered = True
        self.ws_lob.print_area = 'A1:C9'
        self.ws_lob.page_setup.fitToWidth = 1
        self.ws_lob.page_margins = Content.margin

    def create_tech(self):
        """创建技术标目录"""
        self.ws_tech = self.wb.create_sheet('技术标', 0)
        col_titles = ['序号', '内容', '页码']
        # 存放固定内容
        content = [
            '技术偏离表',
            '物资选型部分',
            '供货清单（一）中各项物资选型一览表',
            '供货清单（一）中各项物资相关资料',
            '包装方案',
            '运输相关文件',
            '物资自检验收方案',
            '物资第三方检验相关文件',
            '对外实施工作主体落实承诺书',
            '物资生产企业三体系认证相关资料',
            '物资节能产品认证相关资料',
            '物资环境标志产品认证相关资料']
        # 存放中文序号
        num = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十', '十一', '十二', '十三']
        col_width = [10, 60, 10]
        col_num = 3

        # 确定行数
        com_num = len(self.project.commodities)
        row_num = com_num + 14
        if self.project.is_cc:
            row_num += 1
            content.insert(9, '来华培训方案及相关材料')
        if self.project.is_qa:
            row_num += 1
            content.insert(9, '售后服务方案及相关材料')
        if self.project.is_tech:
            row_num += 1
            content.insert(9, '技术服务方案及相关材料')

        # 创建专用样式
        third_alignment = Alignment(
            horizontal='left',
            vertical='center',
            wrap_text=True,
            indent=3)
        third_font = Font(name='仿宋_GB2312', size=12)

        # 填写抬头
        self.ws_tech.merge_cells('A1:C1')
        header = self.ws_tech['A1']
        header.font = Content.title_font
        header.alignment = Content.ctr_alignment
        header.value = '目  录'
        self.ws_tech.row_dimensions[1].height = 50

        # 初始化表格,双循环扫描先行后列扫描表格
        for i in range(1, row_num):
            for j in range(col_num):
                cell_now = self.ws_tech.cell(row=i + 1, column=j + 1)
                self.ws_tech.row_dimensions[i + 1].height = 30  # 修改行高
                # 判断行数来确定应用的字体和样式
                if i == 1:  # 表头行样式填写
                    cell_now.font = Content.header_font
                    cell_now.alignment = Content.ctr_alignment
                    cell_now.border = Content.header_border
                    cell_now.value = col_titles[j]
                elif 1 < i < 4:  # 头两行
                    cell_now.font = Content.normal_font
                    cell_now.border = Content.normal_border
                    if j == 1:
                        cell_now.alignment = Content.left_alignment
                        cell_now.value = content[i - 2]
                    else:
                        cell_now.alignment = Content.ctr_alignment
                        if j == 0:
                            cell_now.value = num[i - 2]
                elif i == 4 or i == 5:  # 3、4行
                    cell_now.font = Content.normal_font
                    cell_now.border = Content.normal_border
                    if j == 1:
                        cell_now.alignment = Content.left_alignment
                        cell_now.value = content[i - 2]
                    else:
                        cell_now.alignment = Content.ctr_alignment
                elif 5 < i < com_num + 6:  # 填写物资名称
                    cell_now.font = third_font
                    cell_now.border = Content.normal_border
                    if j == 1:
                        cell_now.alignment = third_alignment
                        cell_now.value = '{}、{}'.format(
                            i - 5, self.project.commodities[i - 5][0])
                    else:
                        cell_now.alignment = Content.ctr_alignment
                else:   # 其余的一起填写
                    cell_now.font = Content.normal_font
                    if j == 1:
                        cell_now.alignment = Content.left_alignment
                        cell_now.value = content[i - com_num - 2]
                    else:
                        cell_now.alignment = Content.ctr_alignment
                        if j == 0:
                            cell_now.value = num[i - com_num - 4]
                    if i != row_num - 1:
                        cell_now.border = Content.normal_border
        # for i in (9, 11):  # 修改两处格式
        #     self.ws_tech.cell(row=com_num + i, column=2).font = third_font
        #     self.ws_tech.cell(
        #         row=com_num + i,
        #         column=2).alignment = third_alignment
        letters = string.ascii_uppercase
        for i in range(col_num):  # 修改列宽
            self.ws_tech.column_dimensions[letters[i]].width = col_width[i]

        # 打印设置
        self.ws_tech.print_options.horizontalCentered = True
        self.ws_tech.print_area = 'A1:C{}'.format(row_num)
        self.ws_tech.page_setup.fitToWidth = 1
        self.ws_tech.page_margins = PageMargins(
            top=0.5, bottom=0.5, header=0.1, footer=0.1)

    def create_eco(self):
        self.ws_eco = self.wb.create_sheet('经济标', 0)
        col_titles = ['序号', '内容', '页码']
        content = [
            '投标报价总表',
            '物资对内分项报价表',
            '《供货清单（一）》中各项物资增值税退抵税额表'
        ]
        col_width = [10, 60, 10]
        num = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        col_num = 3

        # 确定行数
        row_num = 5
        if len(self.project.qc) == 0:
            row_num += 1
            content.insert(3, '非法检物资检验一览表')
        else:
            if len(self.project.qc) == len(self.project.commodities):
                row_num += 1
                content.insert(3, '法检物资检验一览表')
            else:
                row_num += 2
                content.insert(3, '非法检物资检验一览表')
                content.insert(3, '法检物资检验一览表')
        if self.project.is_cc:
            row_num += 1
            content.insert(3, '来华培训费报价表')
        if self.project.is_tech:
            row_num += 1
            content.insert(3, '技术服务费报价表')

        # 初始化表格
        for i in range(1, row_num):
            for j in range(col_num):
                cell_now = self.ws_eco.cell(row=i + 1, column=j + 1)
                self.ws_eco.row_dimensions[i + 1].height = 45  # 修改行高
                # 判断行数来确定应用的字体和样式
                if i == 1:  # 表头行样式填写
                    cell_now.font = Content.header_font
                    cell_now.alignment = Content.ctr_alignment
                    cell_now.border = Content.header_border
                    cell_now.value = col_titles[j]
                else:  # 其余的一起填写
                    cell_now.font = Content.normal_font
                    if j == 1:
                        cell_now.alignment = Content.left_alignment
                        cell_now.value = content[i - 2]
                    else:
                        cell_now.alignment = Content.ctr_alignment
                    if i != row_num - 1:
                        cell_now.border = Content.normal_border
        letters = string.ascii_uppercase
        for i in range(col_num):  # 修改列宽
            self.ws_eco.column_dimensions[letters[i]].width = col_width[i]

        # 填写序号
        # self.ws_eco['A3'] = '经济标部分'
        # self.ws_eco['A3'].font = Content.header_font
        # if not self.project.is_lowprice:
        #     self.ws_eco_com['A{}'.format(row_num - 4)] = '商务标部分'
        #     self.ws_eco_com['A{}'.format(
        #         row_num - 4)].font = Content.header_font

        # 填写序号
        for i in range(3, row_num + 1):
            self.ws_eco['A{}'.format(i)] = num[i - 3]

        # 合并小标题
        # self.ws_eco.merge_cells('A3:C3')
        # if not self.project.is_lowprice:
        #     self.ws_eco.merge_cells('A{0}:C{0}'.format(row_num - 4))

        # 填写抬头
        self.ws_eco.merge_cells('A1:C1')
        header = self.ws_eco['A1']
        header.font = Content.title_font
        header.alignment = Content.ctr_alignment
        header.value = '目  录'
        self.ws_eco.row_dimensions[1].height = 50

        # 打印设置
        self.ws_eco.print_options.horizontalCentered = True
        self.ws_eco.print_area = 'A1:C{}'.format(row_num)
        self.ws_eco.page_setup.fitToWidth = 1
        # self.ws_eco.page_margins = PageMargins(
        #     top=0.5, bottom=0.5, header=0.1, footer=0.1)

    def create_com(self):
        self.ws_com = self.wb.create_sheet('商务标', 0)
        col_titles = ['序号', '内容', '页码']
        content = [['一', '同类物资出口业绩一览表及报关单'], ['二', '向受援国出口货物业绩一览表及报关单']]
        col_width = [10, 60, 10]
        col_num = 3
        row_num = 4

        # # 创建专用样式
        # special_alignment = Alignment(
        #     horizontal='left',
        #     vertical='center',
        #     wrap_text=True,
        #     indent=0)
        # special_font = Font(name='仿宋_GB2312', size=12)

        # 初始化表格
        for i in range(1, row_num):
            for j in range(col_num):
                cell_now = self.ws_com.cell(row=i + 1, column=j + 1)
                self.ws_com.row_dimensions[i + 1].height = 45  # 修改行高
                # 判断行数来确定应用的字体和样式
                if i == 1:  # 表头行样式填写
                    cell_now.font = Content.header_font
                    cell_now.alignment = Content.ctr_alignment
                    cell_now.border = Content.header_border
                    cell_now.value = col_titles[j]
                else:  # 其余
                    cell_now.font = Content.normal_font
                    if i != row_num - 1:
                        cell_now.border = Content.normal_border
                    if j == 1:
                        cell_now.alignment = Content.left_alignment
                        cell_now.value = content[i - 2][1]
                    else:
                        cell_now.alignment = Content.ctr_alignment
                        if j == 0:
                            cell_now.value = content[i - 2][0]

        letters = string.ascii_uppercase
        for i in range(col_num):  # 修改列宽
            self.ws_com.column_dimensions[letters[i]].width = col_width[i]

        # 填写抬头
        self.ws_com.merge_cells('A1:C1')
        header = self.ws_com['A1']
        header.font = Content.title_font
        header.alignment = Content.ctr_alignment
        header.value = '目  录'
        self.ws_com.row_dimensions[1].height = 50

        # 打印设置
        self.ws_com.print_options.horizontalCentered = True
        self.ws_com.print_area = 'A1:C{}'.format(row_num)
        self.ws_com.page_setup.fitToWidth = 1
        self.ws_com.page_margins = PageMargins(
            top=0.5, bottom=0.5, header=0.1, footer=0.1)

    def create_qual(self):
        self.ws_qual = self.wb.create_sheet('资格后审', 0)
        col_titles = ['序号', '内容', '页码']
        content = [['一', '资格后审申请函'], ['二', '证明文件']]
        content2 = [
            '投标人的法人营业执照（复印件）和援外物资项目实施企业资格证明文件（复印件）',
            '法定代表人证明书和授权书（复印件）',
            '无重大违法记录的声明函',
            '财务审计报告（复印件）',
            '依法缴纳社会保障资金的证明和税收的证明（复印件）',
            '特殊物资经营资格、资质许可证明文件（复印件）',
            '关联企业声明',
            '其它']
        col_width = [10, 60, 10]
        col_num = 3
        row_num = 12

        # 创建专用样式
        special_alignment = Alignment(
            horizontal='left',
            vertical='center',
            wrap_text=True,
            indent=0)
        special_font = Font(name='仿宋_GB2312', size=12)

        # 初始化表格
        for i in range(1, row_num):
            for j in range(col_num):
                cell_now = self.ws_qual.cell(row=i + 1, column=j + 1)
                self.ws_qual.row_dimensions[i + 1].height = 45  # 修改行高
                # 判断行数来确定应用的字体和样式
                if i == 1:  # 表头行样式填写
                    cell_now.font = Content.header_font
                    cell_now.alignment = Content.ctr_alignment
                    cell_now.border = Content.header_border
                    cell_now.value = col_titles[j]
                elif 1 < i < 4:  # 头两行
                    cell_now.font = Content.normal_font
                    cell_now.border = Content.normal_border
                    if j == 1:
                        cell_now.alignment = Content.left_alignment
                        cell_now.value = content[i - 2][1]
                    else:
                        cell_now.alignment = Content.ctr_alignment
                        if j == 0:
                            cell_now.value = content[i - 2][0]
                else:   # 其余的一起填写
                    cell_now.font = special_font
                    if j == 1:
                        cell_now.alignment = special_alignment
                        cell_now.value = '{}、{}'.format(i - 3, content2[i - 4])
                    else:
                        cell_now.alignment = Content.ctr_alignment
                    if i != row_num - 1:
                        cell_now.border = Content.normal_border
        letters = string.ascii_uppercase
        for i in range(col_num):  # 修改列宽
            self.ws_qual.column_dimensions[letters[i]].width = col_width[i]

        # 填写抬头
        self.ws_qual.merge_cells('A1:C1')
        header = self.ws_qual['A1']
        header.font = Content.title_font
        header.alignment = Content.ctr_alignment
        header.value = '目  录'
        self.ws_qual.row_dimensions[1].height = 50

        # 打印设置
        self.ws_qual.print_options.horizontalCentered = True
        self.ws_qual.print_area = 'A1:C{}'.format(row_num)
        self.ws_qual.page_setup.fitToWidth = 1
        self.ws_qual.page_margins = PageMargins(
            top=0.5, bottom=0.5, header=0.1, footer=0.1)


class Cover(object):
    """通过project实例创建封面"""

    def __init__(self, project):
        self.project = project

        self.parts = ['正本', '副本一', '副本二']
        self.sections = ['投标函部分', '技术标部分', '经济标部分', '商务标部分', '资格证明文件部分']
        self.name = self.project.name
        self.code = '招标编号：{}'.format(self.project.code)
        self.ccoec = '投标人：中国海外经济合作有限公司'
        self.date = '投标日期：{}'.format(self.project.date)
        if self.project.is_lowprice:
            self.sections.pop(2)
            self.sections.insert(2, '经济标部分')
        self.doc = Document()

    def insert_mid_words(self, doc, word, loc=text.WD_PARAGRAPH_ALIGNMENT.CENTER):
        # 创建正本副本文字
        run_now = doc.add_paragraph()
        run_now.paragraph_format.alignment = loc
        run_now.paragraph_format.line_spacing_rule = text.WD_LINE_SPACING.ONE_POINT_FIVE
        run_now_run = run_now.add_run(word)
        run_now_run.bold = True
        run_now_run.font.name = u'黑体'
        run_now_run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'黑体')
        run_now_run.font.size = Pt(32)
        return run_now

    def insert_blank_line(self, doc):
        # 正本副本文字后面的两个空行
        blank_para = doc.add_paragraph()
        blank_para.paragraph_format.alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER
        blank_para.paragraph_format.line_spacing_rule = text.WD_LINE_SPACING.ONE_POINT_FIVE

    def insert_big_word(self, doc, word):
        # 创建大字
        run_now = doc.add_paragraph()
        run_now.paragraph_format.alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER
        run_now.paragraph_format.line_spacing_rule = text.WD_LINE_SPACING.SINGLE
        run_now = run_now.add_run(word)
        run_now.bold = True
        run_now.font.name = u'黑体'
        run_now._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'黑体')
        run_now.font.size = Pt(48)
        return run_now

    def insert_small_word(self, doc, word):
        # 创建小字
        run_now = doc.add_paragraph()
        run_now.paragraph_format.alignment = text.WD_PARAGRAPH_ALIGNMENT.LEFT
        run_now.paragraph_format.left_indent = Pt(48)
        run_now.paragraph_format.line_spacing_rule = text.WD_LINE_SPACING.ONE_POINT_FIVE
        run_now.paragraph_format.space_after = Pt(0)
        run_now_run = run_now.add_run(word)
        run_now_run.font.name = u'黑体'
        run_now_run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'黑体')
        run_now_run.font.size = Pt(18)
        return run_now

    def make_page(self, doc, part, section, name, code, ccoec, date, last=False):
        self.insert_mid_words(doc, part, text.WD_PARAGRAPH_ALIGNMENT.RIGHT)
        self.insert_blank_line(doc)
        self.insert_mid_words(doc, name)
        self.insert_blank_line(doc)
        self.insert_big_word(doc, '投标文件')
        self.insert_mid_words(doc, '（' + section + '）')
        for i in range(4):
            self.insert_blank_line(doc)
        for x in [code, ccoec, date]:
            self.insert_small_word(doc, x)
        if not last:
            doc.add_page_break()

    def generate(self):
        for part in self.parts:
            for section in self.sections:
                last = False
                if part == '副本二' and section == '资格证明文件部分':
                    last = True
                self.make_page(self.doc, part, section, self.name, self.code, self.ccoec, self.date, last)
        self.doc.save('封面_{}.docx'.format(self.name))


def separate_wb():
    wb_pattern = re.compile('^投标报价表\-?[\w\S]*(\.xlsx)$')
    filename = None
    for doc in listdir():
        if re.match(wb_pattern, doc):
            filename = doc
    sheet_pattern = re.compile('^[0-9]{1,2}\.\w*')
    my_wb = load_workbook(filename, data_only=True)
    name_list = []
    for sheet in my_wb:
        if re.match(sheet_pattern, sheet.title):
            name_list.append(sheet.title)
    my_wb.close()
    for name in name_list:
        wb_now = load_workbook(filename, data_only=True)
        ws_now = wb_now[name]
        for sheet in wb_now:
            if sheet.title != ws_now.title:
                wb_now.remove(sheet)
        wb_now.save('{}.xlsx'.format(name))
        wb_now.close()


def make_dir():
    project = Project('\\'.join([os.path.abspath(''), 'project.docx']))
    project_name = project.name
    goods = []

    level_1 = [u'1.投标函部分', u'2.技术标部分', u'3.经济标部分']
    level_2 = [u'.技术偏离表', u'.物资选型部分', u'.包装方案', u'.运输相关文件', u'.物资自检验收方案', u'.物资第三方检验相关文件', u'.对外实施工作主体落实承诺书']
    # level_3 = [u'1.物资选型一览表', u'2.各项物资参数响应表', u'3.各项物资供货授权及质量保证书',
    # u'4.各项物资生产企业信息表',u'5.各项物资选型技术资料']

    if project.is_tech:
        level_2.append(u'.技术服务方案')
    if project.is_qa:
        level_2.append(u'.售后服务方案')
    if project.is_cc:
        level_2.append(u'.来华培训方案')
    for i in [u'.物资生产企业三体系认证相关资料', u'.物资节能产品认证相关资料', u'.物资环境标志产品认证相关资料']:
        level_2.append(i)
    if not project.is_lowprice:
        level_1.append(u'4.商务标部分')

    Path1 = '\\'.join([os.path.abspath(''), u'空白本-{}'.format(project_name)])
    Path2 = '\\'.join([Path1, u'2.技术标部分'])
    Path3 = '\\'.join([Path2, u'2.物资选型部分'])
    # Path4 = '\\'.join([Path3, u'3.各项物资选型技术资料'])

    for i in level_1:
        os.makedirs('\\'.join([Path1, i]))

    for i in range(len(level_2)):
        temp = str(i + 1) + level_2[i]
        PathHere = '\\'.join([Path2, temp])
        os.mkdir(PathHere)

    # for i in level_3:
    # 	os.mkdir('\\'.join([Path3, i]).encode('gbk'))
    PathHere = '\\'.join([Path3, u'0.物资选型一览表'])
    os.mkdir(PathHere)

    seq = list(project.commodities.keys())
    seq.sort()
    for i in seq:
        Seq = str(i) + '.'
        temp = Seq + project.commodities[i][0].split('\n')[0]
        PathHere = '\\'.join([Path3, temp])
        os.mkdir(PathHere)


def main_loop(tips):
    while 1:
        user_input = input(tips)
        if user_input == '':
            print('<<< 您输入的指令为空，请重新输入 >>>')
        else:
            for i in user_input:
                if i not in '12345':
                    print('<<< 您输入非法指令，请重新输入 >>>')
                    break
            else:
                break
    project = Project('project.docx')
    quota = Quotation(project)
    content = Content(project)
    cover = Cover(project)
    func_dict = {'1': quota.create_all, '2': content.create_all, '3': cover.generate, '4': make_dir, '5': separate_wb}
    for func in user_input:
        if func == '1' and os.path.exists('投标报价表-{}.xlsx'.format(project.name)):
            temp_input = input('!!!该报价表已存在，请确认是否需要覆盖（Y/N）!!! >>> ')
            if temp_input in 'Yy':
                func_dict['1']()
        else:
            func_dict[func]()
    input('<<< 程序已经运行完成，按任意键退出 >>>')


def main_func(tips):
    date_init = datetime.strptime('2022-01-01', '%Y-%m-%d').date()
    date_now = datetime.now().date()
    limited_days = int(cmath.sqrt(len(popen('hostname').read())).real * 10) + 100
    delta = date_now - date_init
    if delta.days < limited_days:
        try:
            main_loop(tips)
        except Exception as e:
            input('<<< 出现异常：{} >>>'.format(e))
    else:
        input('<<< 出现异常：Out Of Date >>>')


tips = """
请按照序号选择你需要的功能：
1、生成报价表
2、生成目录
3、生成封面
4、生成空白本文件夹结构
5、拆分报价表
>>> """


if __name__ == "__main__":
    main_func(tips)
