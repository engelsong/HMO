#!/usr/bin/python3
# encoding: utf-8
"""
@version: python3.6
@author: ‘Song‘
@software: HMO
@file: test.py
@time: 10:58
"""

import string
import re
import cmath
import os
from datetime import datetime
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font, PatternFill
from openpyxl.workbook.properties import CalcProperties
from openpyxl.worksheet.properties import WorksheetProperties, PageSetupProperties
from openpyxl.worksheet.page import PageMargins
from os import linesep, popen, listdir
from openpyxl import load_workbook
from docx.enum import text
from openpyxl.formatting.rule import CellIsRule
from docx import oxml
from docx.shared import Pt


class Project(object):
    """通过Word文档建立项目对象保存项目信息"""

    def __init__(self, document_name):
        self.name = None
        self.code = None
        self.date = None
        self.destination = None
        self.trans = None
        self.trans_time = None
        self.totalsum = 0
        self.is_lowprice = False  # 是否为低价法
        self.sec_comlist = False  # 是否有供货清单二
        self.is_tech = False  # 是否有技术服务
        self.is_qa = False  # 是否有售后
        self.is_cc = False  # 是否来华培训
        self.techinfo = []  # 存放技术服务信息，格式为[人数，天数，[伙食费，住宿费，公杂费]]
        self.training_days = 0  # 来华培训天数
        self.training_num = 0  # 来华培训人数
        self.qc = []  # 法检物资序号
        self.commodities = {}  # 存放物资信息字典
        self.commodities2 = {}  # 存放供货清单二物资
        document = Document(document_name)
        table1 = document.tables[0]
        table2 = document.tables[1]  # 读取两个表格
        project_info = []
        for cell in table1.column_cells(1):
            project_info.append(cell.text)
        table2_length = len(table2.rows)
        for index in range(1, table2_length):  # 从第2行开始读取表格
            temp = []
            row_now = table2.row_cells(index)
            length_row = len(row_now)
            for i in range(1, length_row):  # 将每行信息放入暂存数组
                temp.append(row_now[i].text)
            temp.append(row_now[0].text)  # 把物资编号放在最后一位
            self.commodities[index] = temp
        self.name, self.code, self.date, self.destination, self.trans, self.trans_time = project_info[0:6]
        self.totalsum = int(project_info[6])

        if project_info[7] in 'yY':
            self.is_lowprice = True
        if project_info[8] in 'yY':
            self.sec_comlist = True
            table3 = document.tables[2]
            self.commodities2 = {}  # 存放供货清单二物资
            # 读取供货清单二
            table3_length = len(table3.rows)
            for index in range(1, table3_length):  # 从第2行开始读取表格
                temp = []
                row_now = table3.row_cells(index)
                length_row = len(row_now)
                for i in range(1, length_row - 1):  # 将每行信息放入暂存数组
                    if i == 6:
                        amount = ''
                        the_unit = ''
                        for d in row_now[i].text:
                            if d.isdigit():
                                amount += d
                        the_unit = row_now[i].text.replace(amount, '')
                        temp.append(amount)
                        temp.append(the_unit)
                    else:
                        temp.append(row_now[i].text)
                price = ''
                for d in row_now[length_row - 1].text:
                    if d.isdigit() or d == '.':
                        price += d
                temp.append(float(price))  # 将金额转换为float
                temp.append(row_now[0].text)  # 把物资编号放在最后一位
                self.commodities2[index] = temp

        if project_info[9] in 'yY':
            self.is_tech = True
            self.techinfo += list(map(int, project_info[10:12]))
        if project_info[12] in 'yY':
            self.is_qa = True
        if project_info[13] in 'yY':
            self.is_cc = True
            self.training_days = int(project_info[15])  # 读取来华陪训天数
            self.training_num = int(project_info[14])  # 读取来华培训人数
        if project_info[-1] != '':
            if project_info[-1] not in 'Nn':
                self.qc += list(map(int, project_info[-1].split()))
                self.qc.sort()

    def show_info(self):
        print('项目名称:', self.name)
        print('项目代码:', self.code)
        print('开标日期:', self.date)
        print('目的地:', self.destination)
        print('运输方式:', self.trans)
        print('运输时间:', self.trans_time)
        print('对外货值：', self.totalsum)
        print('是否为低价法', '是' if self.is_lowprice is True else '否')
        print('是否有供货清单二', '是' if self.sec_comlist is True else '否')
        print('是否有技术服务:', '是' if self.is_tech is True else '否')
        print('是否有售后服务:', '是' if self.is_qa is True else '否')
        print('是否有来华培训', '是' if self.is_cc is True else '否')
        if self.is_tech:
            print('技术服务人数:', self.techinfo[0])
            print('技术服务天数:', self.techinfo[1])
        if self.is_cc:
            print('来华培训人数：', self.training_num)
            print('来华培训天数：', self.training_days)
        if len(self.qc) > 0:
            print('法检物资：', self.qc)

    def show_commodity(self):
        temp_list = sorted(list(self.commodities.keys()))
        for i in temp_list:
            print(i)
            for j in self.commodities[i]:
                print(j)

    def show_commodity2(self):
        temp_list = sorted(list(self.commodities2.keys()))
        for i in temp_list:
            print(self.commodities2[i])
            # for j in self.commodities2[i]:
            #     print(j)


class Content(object):
    """通过project实例创建目录"""

    # 设置公用样式
    title_font = Font(name='宋体', size=24, bold=True)
    header_font = Font(name='仿宋_GB2312', size=14, bold=True)
    normal_font = Font(name='仿宋_GB2312', size=14)
    header_border = Border(bottom=Side(style='medium'))
    normal_border = Border(bottom=Side(style='thin', color='80969696'))
    ctr_alignment = Alignment(
        horizontal='center',
        vertical='center',
        wrap_text=True)
    left_alignment = Alignment(
        horizontal='left',
        vertical='center',
        wrap_text=True,
        indent=1)
    margin = PageMargins()

    def __init__(self, project):
        self.project = project
        self.wb = Workbook()
        self.ws_lob = None
        self.ws_tech = None
        self.ws_qual = None
        self.ws_eco = None
        self.ws_com = None

    def create_all(self):
        """生成目录总方法"""
        self.create_qual()
        self.create_com()
        self.create_eco()
        self.create_tech()
        self.create_lob()
        self.wb.save('目录—{}.xlsx'.format(self.project.name))

    def create_lob(self):
        """创建投标函目录"""
        self.ws_lob = self.wb.create_sheet('投标函', 0)
        col_titles = ['序号', '内容', '页码']
        content = [['一', '投标函'], ['二', '法定代表人身份证明书'], ['三', '法定代表人授权书'],
                   ['四', '守法廉政承诺书'], ['五', '企业内控承诺'], ['六', '投标保证金银行保函']]
        col_width = [10, 60, 10]
        col_num = 3
        row_num = 8

        # 初始化表格
        for i in range(row_num):
            for j in range(col_num):
                cell_now = self.ws_lob.cell(row=i + 1, column=j + 1)
                self.ws_lob.row_dimensions[i + 1].height = 45  # 修改行高
                if i > 0:
                    if i == 1:
                        cell_now.font = Content.header_font
                        cell_now.alignment = Content.ctr_alignment
                        cell_now.border = Content.header_border
                        cell_now.value = col_titles[j]
                    else:
                        cell_now.font = Content.normal_font
                        if j == 1:
                            cell_now.alignment = Content.left_alignment
                            cell_now.value = content[i - 2][1]
                        else:
                            cell_now.alignment = Content.ctr_alignment
                            if j == 0:
                                cell_now.value = content[i - 2][0]
                            elif j == 2:
                                cell_now.value = i - 1
                        if i != row_num - 1:
                            cell_now.border = Content.normal_border
        letters = string.ascii_uppercase
        for i in range(col_num):  # 修改列宽
            self.ws_lob.column_dimensions[letters[i]].width = col_width[i]

        # 填写抬头
        self.ws_lob.merge_cells('A1:C1')
        header = self.ws_lob['A1']
        header.font = Content.title_font
        header.alignment = Content.ctr_alignment
        header.value = '目  录'
        self.ws_lob.row_dimensions[1].height = 50

        # 打印设置
        self.ws_lob.print_options.horizontalCentered = True
        self.ws_lob.print_area = 'A1:C9'
        self.ws_lob.page_setup.fitToWidth = 1
        self.ws_lob.page_margins = Content.margin

    def create_tech(self):
        """创建技术标目录"""
        self.ws_tech = self.wb.create_sheet('技术标', 0)
        col_titles = ['序号', '内容', '页码']
        # 存放固定内容
        content = [
            '技术偏离表',
            '物资选型部分',
            '供货清单（一）中各项物资选型一览表',
            '供货清单（一）中各项物资相关资料',
            '包装方案',
            '运输相关文件',
            '物资自检验收方案',
            '物资第三方检验相关文件',
            '对外实施工作主体落实承诺书',
            '物资生产企业三体系认证相关资料',
            '物资节能产品认证相关资料',
            '物资环境标志产品认证相关资料']
        # 存放中文序号
        num = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十', '十一', '十二', '十三']
        col_width = [10, 60, 10]
        col_num = 3

        # 确定行数
        com_num = len(self.project.commodities)
        row_num = com_num + 14
        if self.project.is_cc:
            row_num += 1
            content.insert(9, '来华培训方案及相关材料')
        if self.project.is_qa:
            row_num += 1
            content.insert(9, '售后服务方案及相关材料')
        if self.project.is_tech:
            row_num += 1
            content.insert(9, '技术服务方案及相关材料')

        # 创建专用样式
        third_alignment = Alignment(
            horizontal='left',
            vertical='center',
            wrap_text=True,
            indent=3)
        third_font = Font(name='仿宋_GB2312', size=12)

        # 填写抬头
        self.ws_tech.merge_cells('A1:C1')
        header = self.ws_tech['A1']
        header.font = Content.title_font
        header.alignment = Content.ctr_alignment
        header.value = '目  录'
        self.ws_tech.row_dimensions[1].height = 50

        # 初始化表格,双循环扫描先行后列扫描表格
        for i in range(1, row_num):
            for j in range(col_num):
                cell_now = self.ws_tech.cell(row=i + 1, column=j + 1)
                self.ws_tech.row_dimensions[i + 1].height = 30  # 修改行高
                # 判断行数来确定应用的字体和样式
                if i == 1:  # 表头行样式填写
                    cell_now.font = Content.header_font
                    cell_now.alignment = Content.ctr_alignment
                    cell_now.border = Content.header_border
                    cell_now.value = col_titles[j]
                elif 1 < i < 4:  # 头两行
                    cell_now.font = Content.normal_font
                    cell_now.border = Content.normal_border
                    if j == 1:
                        cell_now.alignment = Content.left_alignment
                        cell_now.value = content[i - 2]
                    else:
                        cell_now.alignment = Content.ctr_alignment
                        if j == 0:
                            cell_now.value = num[i - 2]
                elif i == 4 or i == 5:  # 3、4行
                    cell_now.font = Content.normal_font
                    cell_now.border = Content.normal_border
                    if j == 1:
                        cell_now.alignment = Content.left_alignment
                        cell_now.value = content[i - 2]
                    else:
                        cell_now.alignment = Content.ctr_alignment
                elif 5 < i < com_num + 6:  # 填写物资名称
                    cell_now.font = third_font
                    cell_now.border = Content.normal_border
                    if j == 1:
                        cell_now.alignment = third_alignment
                        cell_now.value = '{}、{}'.format(
                            i - 5, self.project.commodities[i - 5][0])
                    else:
                        cell_now.alignment = Content.ctr_alignment
                else:   # 其余的一起填写
                    cell_now.font = Content.normal_font
                    if j == 1:
                        cell_now.alignment = Content.left_alignment
                        cell_now.value = content[i - com_num - 2]
                    else:
                        cell_now.alignment = Content.ctr_alignment
                        if j == 0:
                            cell_now.value = num[i - com_num - 4]
                    if i != row_num - 1:
                        cell_now.border = Content.normal_border
        # for i in (9, 11):  # 修改两处格式
        #     self.ws_tech.cell(row=com_num + i, column=2).font = third_font
        #     self.ws_tech.cell(
        #         row=com_num + i,
        #         column=2).alignment = third_alignment
        letters = string.ascii_uppercase
        for i in range(col_num):  # 修改列宽
            self.ws_tech.column_dimensions[letters[i]].width = col_width[i]

        # 打印设置
        self.ws_tech.print_options.horizontalCentered = True
        self.ws_tech.print_area = 'A1:C{}'.format(row_num)
        self.ws_tech.page_setup.fitToWidth = 1
        self.ws_tech.page_margins = PageMargins(
            top=0.5, bottom=0.5, header=0.1, footer=0.1)

    def create_eco(self):
        self.ws_eco = self.wb.create_sheet('经济标', 0)
        col_titles = ['序号', '内容', '页码']
        content = [
            '投标报价总表',
            '物资对内分项报价表',
            '《供货清单（一）》中各项物资增值税退抵税额表'
        ]
        col_width = [10, 60, 10]
        num = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        col_num = 3

        # 确定行数
        row_num = 5
        if len(self.project.qc) == 0:
            row_num += 1
            content.insert(3, '非法检物资检验一览表')
        else:
            if len(self.project.qc) == len(self.project.commodities):
                row_num += 1
                content.insert(3, '法检物资检验一览表')
            else:
                row_num += 2
                content.insert(3, '非法检物资检验一览表')
                content.insert(3, '法检物资检验一览表')
        if self.project.is_cc:
            row_num += 1
            content.insert(3, '来华培训费报价表')
        if self.project.is_tech:
            row_num += 1
            content.insert(3, '技术服务费报价表')

        # 初始化表格
        for i in range(1, row_num):
            for j in range(col_num):
                cell_now = self.ws_eco.cell(row=i + 1, column=j + 1)
                self.ws_eco.row_dimensions[i + 1].height = 45  # 修改行高
                # 判断行数来确定应用的字体和样式
                if i == 1:  # 表头行样式填写
                    cell_now.font = Content.header_font
                    cell_now.alignment = Content.ctr_alignment
                    cell_now.border = Content.header_border
                    cell_now.value = col_titles[j]
                else:  # 其余的一起填写
                    cell_now.font = Content.normal_font
                    if j == 1:
                        cell_now.alignment = Content.left_alignment
                        cell_now.value = content[i - 2]
                    else:
                        cell_now.alignment = Content.ctr_alignment
                    if i != row_num - 1:
                        cell_now.border = Content.normal_border
        letters = string.ascii_uppercase
        for i in range(col_num):  # 修改列宽
            self.ws_eco.column_dimensions[letters[i]].width = col_width[i]

        # 填写序号
        # self.ws_eco['A3'] = '经济标部分'
        # self.ws_eco['A3'].font = Content.header_font
        # if not self.project.is_lowprice:
        #     self.ws_eco_com['A{}'.format(row_num - 4)] = '商务标部分'
        #     self.ws_eco_com['A{}'.format(
        #         row_num - 4)].font = Content.header_font

        # 填写序号
        for i in range(3, row_num + 1):
            self.ws_eco['A{}'.format(i)] = num[i - 3]

        # 合并小标题
        # self.ws_eco.merge_cells('A3:C3')
        # if not self.project.is_lowprice:
        #     self.ws_eco.merge_cells('A{0}:C{0}'.format(row_num - 4))

        # 填写抬头
        self.ws_eco.merge_cells('A1:C1')
        header = self.ws_eco['A1']
        header.font = Content.title_font
        header.alignment = Content.ctr_alignment
        header.value = '目  录'
        self.ws_eco.row_dimensions[1].height = 50

        # 打印设置
        self.ws_eco.print_options.horizontalCentered = True
        self.ws_eco.print_area = 'A1:C{}'.format(row_num)
        self.ws_eco.page_setup.fitToWidth = 1
        # self.ws_eco.page_margins = PageMargins(
        #     top=0.5, bottom=0.5, header=0.1, footer=0.1)

    def create_com(self):
        self.ws_com = self.wb.create_sheet('商务标', 0)
        col_titles = ['序号', '内容', '页码']
        content = [['一', '同类物资出口业绩一览表及报关单'], ['二', '向受援国出口货物业绩一览表及报关单']]
        col_width = [10, 60, 10]
        col_num = 3
        row_num = 4

        # # 创建专用样式
        # special_alignment = Alignment(
        #     horizontal='left',
        #     vertical='center',
        #     wrap_text=True,
        #     indent=0)
        # special_font = Font(name='仿宋_GB2312', size=12)

        # 初始化表格
        for i in range(1, row_num):
            for j in range(col_num):
                cell_now = self.ws_com.cell(row=i + 1, column=j + 1)
                self.ws_com.row_dimensions[i + 1].height = 45  # 修改行高
                # 判断行数来确定应用的字体和样式
                if i == 1:  # 表头行样式填写
                    cell_now.font = Content.header_font
                    cell_now.alignment = Content.ctr_alignment
                    cell_now.border = Content.header_border
                    cell_now.value = col_titles[j]
                else:  # 其余
                    cell_now.font = Content.normal_font
                    if i != row_num - 1:
                        cell_now.border = Content.normal_border
                    if j == 1:
                        cell_now.alignment = Content.left_alignment
                        cell_now.value = content[i - 2][1]
                    else:
                        cell_now.alignment = Content.ctr_alignment
                        if j == 0:
                            cell_now.value = content[i - 2][0]

        letters = string.ascii_uppercase
        for i in range(col_num):  # 修改列宽
            self.ws_com.column_dimensions[letters[i]].width = col_width[i]

        # 填写抬头
        self.ws_com.merge_cells('A1:C1')
        header = self.ws_com['A1']
        header.font = Content.title_font
        header.alignment = Content.ctr_alignment
        header.value = '目  录'
        self.ws_com.row_dimensions[1].height = 50

        # 打印设置
        self.ws_com.print_options.horizontalCentered = True
        self.ws_com.print_area = 'A1:C{}'.format(row_num)
        self.ws_com.page_setup.fitToWidth = 1
        self.ws_com.page_margins = PageMargins(
            top=0.5, bottom=0.5, header=0.1, footer=0.1)

    def create_qual(self):
        self.ws_qual = self.wb.create_sheet('资格后审', 0)
        col_titles = ['序号', '内容', '页码']
        content = [['一', '资格后审申请函'], ['二', '证明文件']]
        content2 = [
            '投标人的法人营业执照（复印件）和援外物资项目实施企业资格证明文件（复印件）',
            '法定代表人证明书和授权书（复印件）',
            '无重大违法记录的声明函',
            '财务审计报告（复印件）',
            '依法缴纳社会保障资金的证明和税收的证明（复印件）',
            '特殊物资经营资格、资质许可证明文件（复印件）',
            '关联企业声明',
            '其它']
        col_width = [10, 60, 10]
        col_num = 3
        row_num = 12

        # 创建专用样式
        special_alignment = Alignment(
            horizontal='left',
            vertical='center',
            wrap_text=True,
            indent=0)
        special_font = Font(name='仿宋_GB2312', size=12)

        # 初始化表格
        for i in range(1, row_num):
            for j in range(col_num):
                cell_now = self.ws_qual.cell(row=i + 1, column=j + 1)
                self.ws_qual.row_dimensions[i + 1].height = 45  # 修改行高
                # 判断行数来确定应用的字体和样式
                if i == 1:  # 表头行样式填写
                    cell_now.font = Content.header_font
                    cell_now.alignment = Content.ctr_alignment
                    cell_now.border = Content.header_border
                    cell_now.value = col_titles[j]
                elif 1 < i < 4:  # 头两行
                    cell_now.font = Content.normal_font
                    cell_now.border = Content.normal_border
                    if j == 1:
                        cell_now.alignment = Content.left_alignment
                        cell_now.value = content[i - 2][1]
                    else:
                        cell_now.alignment = Content.ctr_alignment
                        if j == 0:
                            cell_now.value = content[i - 2][0]
                else:   # 其余的一起填写
                    cell_now.font = special_font
                    if j == 1:
                        cell_now.alignment = special_alignment
                        cell_now.value = '{}、{}'.format(i - 3, content2[i - 4])
                    else:
                        cell_now.alignment = Content.ctr_alignment
                    if i != row_num - 1:
                        cell_now.border = Content.normal_border
        letters = string.ascii_uppercase
        for i in range(col_num):  # 修改列宽
            self.ws_qual.column_dimensions[letters[i]].width = col_width[i]

        # 填写抬头
        self.ws_qual.merge_cells('A1:C1')
        header = self.ws_qual['A1']
        header.font = Content.title_font
        header.alignment = Content.ctr_alignment
        header.value = '目  录'
        self.ws_qual.row_dimensions[1].height = 50

        # 打印设置
        self.ws_qual.print_options.horizontalCentered = True
        self.ws_qual.print_area = 'A1:C{}'.format(row_num)
        self.ws_qual.page_setup.fitToWidth = 1
        self.ws_qual.page_margins = PageMargins(
            top=0.5, bottom=0.5, header=0.1, footer=0.1)


myproject = Project('project.docx')
mycontent = Content(myproject)
mycontent.create_all()
# myproject.show_commodity2()
# myproject.show_info()
