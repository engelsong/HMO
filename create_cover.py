# coding: UTF-8

from docx import Document
from docx.enum import text
from docx.shared import Pt
from docx import oxml
import cmath
from os import popen
from datetime import datetime


class Project(object):
    """通过Word文档建立项目对象保存项目信息"""

    def __init__(self, document_name):
        self.name = None
        self.code = None
        self.date = None
        self.destination = None
        self.trans = None
        self.totalsum = 0
        self.is_lowprice = False  # 是否为低价法
        self.is_tech = False  # 是否有技术服务
        self.is_qa = False  # 是否有售后
        self.is_cc = False  # 是否来华培训
        self.techinfo = []  # 存放技术服务信息，格式为[人数，天数，[伙食费，住宿费，公杂费]]
        self.training_days = 0  # 来华培训天数
        self.training_num = 0  # 来华培训人数
        self.qc = []  # 法检物资序号
        self.commodities = {}  # 存放物资信息字典
        document = Document(document_name)
        table1, table2 = document.tables  # 读取两个表格
        project_info = []
        for cell in table1.column_cells(1):
            project_info.append(cell.text)
        table2_length = len(table2.rows)
        for index in range(1, table2_length):  # 从第2行开始读取表格
            temp = []
            row_now = table2.row_cells(index)
            length_row = len(row_now)
            for i in range(1, length_row):  # 将每行信息放入暂存数组
                temp.append(row_now[i].text)
            temp.append(row_now[0].text)  # 把物资编号放在最后一位
            self.commodities[index] = temp
        self.name, self.code, self.date, self.destination, self.trans = project_info[0:5]
        self.totalsum = int(project_info[5])
        if project_info[6] in 'yY':
            self.is_lowprice = True
        if project_info[7] in 'yY':
            self.is_tech = True
            self.techinfo += list(map(int, project_info[8:10]))
            self.techinfo.append(list(map(int, project_info[10].split())))
        if project_info[11] in 'yY':
            self.is_qa = True
        if project_info[12] in 'yY':
            self.is_cc = True
            self.training_days = int(project_info[14])  # 读取来华陪训天数
            self.training_num = int(project_info[13])  # 读取来华培训人数
        if project_info[-1] != '':
            self.qc += list(map(int, project_info[-1].split()))
            self.qc.sort()

    def show_info(self):
        print('项目名称:', self.name)
        print('项目代码:', self.code)
        print('开标日期:', self.date)
        print('目的地:', self.destination)
        print('运输方式:', self.trans)
        print('对外货值：', self.totalsum)
        print('是否为低价法', '是' if self.is_lowprice is True else '否')
        print('是否有技术服务:', '是' if self.is_tech is True else '否')
        print('是否有售后服务:', '是' if self.is_qa is True else '否')
        print('是否有来华培训', '是' if self.is_cc is True else '否')
        if self.is_tech:
            print('技术服务人数:', self.techinfo[0])
            print('技术服务天数:', self.techinfo[1])
            print('伙食费:', self.techinfo[2][0])
            print('住宿费:', self.techinfo[2][1])
            print('公杂费:', self.techinfo[2][2])
        if self.is_cc:
            print('来华培训人数：', self.training_num)
            print('来华培训天数：', self.training_days)
        if len(self.qc) > 0:
            print('法检物资：', self.qc)

    def show_commoditiy(self):
        temp_list = sorted(list(self.commodities.keys()))
        for i in temp_list:
            print(i)
            for j in self.commodities[i]:
                print(j)


class Cover(object):
    """通过project实例创建封面"""
    # parts = ['正本', '副本一', '副本二']
    # sections = ['投标函部分', '技术标部分', '经济标和商务标部分', '资格证明文件部分']
    # name = self.project.name
    # code = '招标编号：{}'.format(self.project.code)
    # ccoec = '投标人：中国海外经济合作有限公司'
    # date = '投标日期：{}'.format(self.project.date)

    def __init__(self, project):
        self.project = project

        self.parts = ['正本', '副本一', '副本二']
        self.sections = ['投标函部分', '技术标部分', '经济标和商务标部分', '资格证明文件部分']
        self.name = self.project.name
        self.code = '招标编号：{}'.format(self.project.code)
        self.ccoec = '投标人：中国海外经济合作有限公司'
        self.date = '投标日期：{}'.format(self.project.date)
        if self.project.is_lowprice:
            self.sections.pop(2)
            self.sections.insert(2, '经济标部分')
        self.doc = Document()

    def insert_mid_words(self, doc, word, loc=text.WD_PARAGRAPH_ALIGNMENT.CENTER):
        # 创建正本副本文字
        run_now = doc.add_paragraph()
        run_now.paragraph_format.alignment = loc
        run_now.paragraph_format.line_spacing_rule = text.WD_LINE_SPACING.ONE_POINT_FIVE
        run_now_run = run_now.add_run(word)
        run_now_run.bold = True
        run_now_run.font.name = u'黑体'
        run_now_run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'黑体')
        run_now_run.font.size = Pt(32)
        return run_now

    def insert_blank_line(self, doc):
        # 正本副本文字后面的两个空行
        blank_para = doc.add_paragraph()
        blank_para.paragraph_format.alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER
        blank_para.paragraph_format.line_spacing_rule = text.WD_LINE_SPACING.ONE_POINT_FIVE

    def insert_big_word(self, doc, word):
        # 创建大字
        run_now = doc.add_paragraph()
        run_now.paragraph_format.alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER
        run_now.paragraph_format.line_spacing_rule = text.WD_LINE_SPACING.SINGLE
        run_now = run_now.add_run(word)
        run_now.bold = True
        run_now.font.name = u'黑体'
        run_now._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'黑体')
        run_now.font.size = Pt(48)
        return run_now

    def insert_small_word(self, doc, word):
        # 创建小字
        run_now = doc.add_paragraph()
        run_now.paragraph_format.alignment = text.WD_PARAGRAPH_ALIGNMENT.LEFT
        run_now.paragraph_format.left_indent = Pt(48)
        run_now.paragraph_format.line_spacing_rule = text.WD_LINE_SPACING.ONE_POINT_FIVE
        run_now.paragraph_format.space_after = Pt(0)
        run_now_run = run_now.add_run(word)
        run_now_run.font.name = u'黑体'
        run_now_run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'黑体')
        run_now_run.font.size = Pt(18)
        return run_now

    def make_page(self, doc, part, section, name, code, ccoec, date, last=False):
        self.insert_mid_words(doc, part, text.WD_PARAGRAPH_ALIGNMENT.RIGHT)
        self.insert_blank_line(doc)
        self.insert_mid_words(doc, name)
        self.insert_blank_line(doc)
        self.insert_big_word(doc, '投标文件')
        self.insert_mid_words(doc, '（' + section + '）')
        for i in range(4):
            self.insert_blank_line(doc)
        for x in [code, ccoec, date]:
            self.insert_small_word(doc, x)
        if not last:
            doc.add_page_break()

    def gener_cover(self):
        for part in self.parts:
            for section in self.sections:
                last = False
                if part == '副本二' and section == '资格证明文件部分':
                    last = True
                self.make_page(self.doc, part, section, self.name, self.code, self.ccoec, self.date, last)
        self.doc.save('封面_{}.docx'.format(self.name))


project = Project('project.docx')
my_cover = Cover(project)
my_cover.gener_cover()





